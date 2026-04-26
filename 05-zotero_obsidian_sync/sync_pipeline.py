from __future__ import annotations

import argparse
import difflib
import hashlib
import html
import json
import mimetypes
import os
import re
import shutil
import sqlite3
import subprocess
import sys
import time
import unicodedata
import urllib.error
import urllib.parse
import urllib.request
import uuid
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from shutil import which


ARXIV_RE = re.compile(r"\b(\d{4}\.\d{4,5})(v\d+)?\b", re.IGNORECASE)
DOI_RE = re.compile(r"\b10\.\d{4,9}/[-._;()/:A-Z0-9]+\b", re.IGNORECASE)
YEAR_RE = re.compile(r"\b(19|20)\d{2}\b")
SUMMARY_NOTE_MARKER = "codex-zotero-summary-note"
SUMMARY_ATTACHMENT_TITLE_PREFIX = "AI Summary (Codex)"
DEEP_PACKAGE_NOTE_MARKER = "codex-zotero-deep-reading-package"
DEEP_PACKAGE_ATTACHMENT_TITLE_PREFIX = "Deep Reading Package (Codex)"
EXTRACTOR_MIN_VERSION = 3


@dataclass
class Config:
    config_path: Path
    workspace_root: Path
    paper_dir: Path
    summary_root: Path
    spec_path: Path
    state_path: Path
    cache_root: Path
    export_path: Path
    poll_interval_seconds: int
    append_auto_materials: bool
    openai_enabled: bool
    openai_api_key_env: str
    openai_model: str
    openai_reasoning_effort: str | None
    openai_max_input_chars: int
    openai_endpoint: str
    compatible_enabled: bool
    compatible_api_key_env: str
    compatible_model: str
    compatible_base_url: str
    compatible_endpoint: str
    compatible_temperature: float
    summary_backend_preference: str
    codex_cli_enabled: bool
    codex_cli_command: str
    codex_cli_model: str | None
    codex_cli_reasoning_effort: str | None
    zotero_api_enabled: bool
    zotero_api_key_env: str
    zotero_user_id: str
    zotero_library_type: str
    zotero_update_summary_note: bool
    zotero_attach_summary_markdown: bool
    zotero_summary_attachment_mode: str
    obsidian_vault_dir: Path
    obsidian_papers_subdir: str
    obsidian_tags_subdir: str
    obsidian_tag_prefix: str
    bbt_jsonrpc_url: str
    bbt_collection_path: str
    bbt_translator: str
    bbt_display_options: dict[str, Any]
    bbt_replace: bool
    zotero_local_enabled: bool
    zotero_local_db_path: Path | None
    zotero_local_collection_name: str
    zotero_local_fallback_to_all_pdf_items: bool
    tools_python: Path
    extract_script: Path
    package_script: Path
    summary_user_request: str


class SyncError(RuntimeError):
    pass


def log(message: str) -> None:
    try:
        print(message, flush=True)
    except UnicodeEncodeError:
        safe = message.encode(sys.stdout.encoding or "utf-8", errors="replace").decode(
            sys.stdout.encoding or "utf-8",
            errors="replace",
        )
        print(safe, flush=True)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Sync Zotero-linked PDFs in 01-paper to AI summaries, packaged folders, and an Obsidian vault."
    )
    parser.add_argument(
        "command",
        nargs="?",
        default="once",
        choices=("once", "compare", "deep-package", "check-summary", "watch", "setup-bbt"),
        help="Run once, compare selected PDFs, generate deep packages, QA summaries, watch, or setup Better BibTeX.",
    )
    parser.add_argument(
        "--config",
        type=Path,
        default=Path(__file__).resolve().parent / "config.json",
        help="Path to config.json.",
    )
    parser.add_argument(
        "--interval",
        type=int,
        help="Override the watch polling interval in seconds.",
    )
    parser.add_argument(
        "--pdf",
        type=Path,
        nargs="*",
        help="Optional explicit PDF paths. Defaults to all PDFs under paper_dir.",
    )
    parser.add_argument(
        "--summary-backend",
        choices=("auto", "openai", "openai_compatible", "codex"),
        help="Override summary backend selection for this run.",
    )
    parser.add_argument(
        "--summary-model",
        help="Override the model used by the selected summary backend for this run.",
    )
    parser.add_argument(
        "--summary-reasoning-effort",
        choices=("low", "medium", "high", "xhigh"),
        help="Override reasoning effort for the selected summary backend for this run.",
    )
    parser.add_argument(
        "--summary-user-request",
        default="",
        help="Extra user requirements and key questions to combine with the default summary spec.",
    )
    parser.add_argument(
        "--force-regenerate",
        action="store_true",
        help="Regenerate summaries even when an existing Markdown summary is available.",
    )
    return parser.parse_args()


def load_json(path: Path, default: Any) -> Any:
    if not path.exists():
        return default
    return json.loads(path.read_text(encoding="utf-8"))


def save_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        while True:
            chunk = handle.read(1024 * 1024)
            if not chunk:
                break
            digest.update(chunk)
    return digest.hexdigest()


def md5_file(path: Path) -> str:
    digest = hashlib.md5()
    with path.open("rb") as handle:
        while True:
            chunk = handle.read(1024 * 1024)
            if not chunk:
                break
            digest.update(chunk)
    return digest.hexdigest()


def file_signature(path: Path) -> dict[str, int]:
    stat = path.stat()
    return {"size": int(stat.st_size), "mtime_ns": int(stat.st_mtime_ns)}


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def first_nonempty(*values: Any) -> str:
    for value in values:
        if isinstance(value, str):
            stripped = value.strip()
            if stripped:
                return stripped
    return ""


def ascii_slug(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", value)
    ascii_value = normalized.encode("ascii", "ignore").decode("ascii")
    ascii_value = re.sub(r"[^A-Za-z0-9._-]+", "-", ascii_value)
    ascii_value = re.sub(r"-{2,}", "-", ascii_value)
    return ascii_value.strip("-") or "item"


def filename_fragment(value: str, max_words: int, max_length: int) -> str:
    normalized = unicodedata.normalize("NFKD", value)
    ascii_value = normalized.encode("ascii", "ignore").decode("ascii")
    words = re.findall(r"[A-Za-z0-9]+", ascii_value)
    fragment = "-".join(words[:max_words])
    fragment = re.sub(r"-{2,}", "-", fragment).strip("-")
    return fragment[:max_length].strip("-")


def sanitize_windows_filename(value: str, default: str, max_length: int = 160) -> str:
    text = first_nonempty(value) or default
    text = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "-", text)
    text = re.sub(r"\s+", " ", text).strip().rstrip(".")
    if not text:
        text = default
    if len(text) > max_length:
        text = text[:max_length].rstrip(" .")
    return text or default


def normalize_title(value: str) -> str:
    lowered = value.lower()
    lowered = re.sub(r"[^a-z0-9\u4e00-\u9fff]+", "", lowered)
    return lowered


def normalize_identifier(value: str) -> str:
    return re.sub(r"\s+", "", value).lower()


def sanitize_obsidian_tag(tag: str) -> str:
    tag = tag.strip()
    tag = re.sub(r"\s+", "-", tag)
    tag = re.sub(r"[^\w/\-\u4e00-\u9fff]+", "-", tag, flags=re.UNICODE)
    tag = re.sub(r"-{2,}", "-", tag)
    return tag.strip("-/") or "untagged"


def yaml_quote(value: str) -> str:
    return json.dumps(value, ensure_ascii=False)


def strip_markdown_fences(text: str) -> str:
    stripped = text.strip()
    if stripped.startswith("```") and stripped.endswith("```"):
        lines = stripped.splitlines()
        if len(lines) >= 2:
            return "\n".join(lines[1:-1]).strip() + "\n"
    return text.strip() + "\n"


def path_to_uri(path: Path) -> str:
    try:
        return path.resolve().as_uri()
    except ValueError:
        text = str(path.resolve()).replace("\\", "/")
        if text.startswith("//"):
            return f"file:{text}"
        return f"file:///{text.lstrip('/')}"


def resolve_existing_path(value: Any, *, kind: str = "any") -> Path | None:
    text = first_nonempty(value)
    if not text:
        return None
    candidate = Path(text)
    try:
        if kind == "file":
            return candidate.resolve() if candidate.is_file() else None
        if kind == "dir":
            return candidate.resolve() if candidate.is_dir() else None
        return candidate.resolve() if candidate.exists() else None
    except OSError:
        return None


def first_existing_path(candidates: list[Path], *, kind: str = "any") -> Path | None:
    for candidate in candidates:
        try:
            if kind == "file":
                if candidate.is_file():
                    return candidate.resolve()
                continue
            if kind == "dir":
                if candidate.is_dir():
                    return candidate.resolve()
                continue
            if candidate.exists():
                return candidate.resolve()
        except OSError:
            continue
    return None


def is_path_under(path: Path, parent: Path) -> bool:
    try:
        path.resolve().relative_to(parent.resolve())
        return True
    except ValueError:
        return False
    except OSError:
        return False


def preferred_record_pdf_path(config: Config, candidate_pdf: object, existing_pdf: object) -> str:
    existing = first_nonempty(existing_pdf)
    candidate = first_nonempty(candidate_pdf)
    if existing:
        try:
            existing_path = Path(existing)
            if existing_path.exists() and is_path_under(existing_path, config.paper_dir):
                return str(existing_path.resolve())
        except OSError:
            pass
    if candidate:
        return candidate
    return existing


def repair_record_paths(config: Config, record: dict[str, Any] | None) -> dict[str, Any] | None:
    if not isinstance(record, dict):
        return record

    repaired = dict(record)
    summary_stem = first_nonempty(repaired.get("summary_stem"))
    obsidian_stem = first_nonempty(repaired.get("obsidian_stem"), summary_stem)
    papers_root = config.obsidian_vault_dir / config.obsidian_papers_subdir

    if summary_stem:
        summary_candidates = []
        summary_value = first_nonempty(repaired.get("summary_md"))
        if summary_value:
            summary_candidates.append(Path(summary_value))
        summary_candidates.append(config.summary_root / f"{summary_stem}.md")
        summary_candidates.append(config.summary_root / summary_stem / f"{summary_stem}.md")
        summary_path = first_existing_path(summary_candidates, kind="file")
        if summary_path is not None:
            repaired["summary_md"] = str(summary_path)

        package_candidates = []
        package_value = first_nonempty(repaired.get("package_dir"))
        if package_value:
            package_candidates.append(Path(package_value))
        package_candidates.append(config.summary_root / summary_stem)
        package_dir = first_existing_path(package_candidates, kind="dir")
        if package_dir is not None:
            repaired["package_dir"] = str(package_dir)

    if obsidian_stem:
        obsidian_dir_candidates = []
        obsidian_dir_value = first_nonempty(repaired.get("obsidian_dir"))
        if obsidian_dir_value:
            obsidian_dir_candidates.append(Path(obsidian_dir_value))
        obsidian_dir_candidates.append(papers_root / obsidian_stem)
        obsidian_dir = first_existing_path(obsidian_dir_candidates, kind="dir")
        if obsidian_dir is not None:
            repaired["obsidian_dir"] = str(obsidian_dir)
            obsidian_md_candidates = []
            obsidian_md_value = first_nonempty(repaired.get("obsidian_md"))
            if obsidian_md_value:
                obsidian_md_candidates.append(Path(obsidian_md_value))
            obsidian_md_candidates.append(obsidian_dir / f"{obsidian_stem}.md")
            obsidian_md = first_existing_path(obsidian_md_candidates, kind="file")
            if obsidian_md is not None:
                repaired["obsidian_md"] = str(obsidian_md)

    return repaired


def zotero_link_cache_root() -> Path:
    local_appdata = os.environ.get("LOCALAPPDATA", "").strip()
    if local_appdata:
        base = Path(local_appdata)
    else:
        base = Path.home() / "AppData" / "Local"
    return base / "PaperSyncManager" / "zotero_link_cache"


def prepare_zotero_linked_summary_path(record: dict[str, Any]) -> Path:
    package_dir = Path(first_nonempty(record.get("package_dir")))
    summary_stem = first_nonempty(record.get("summary_stem"), package_dir.name, "summary")
    if package_dir.exists():
        cache_dir_name = sanitize_windows_filename(summary_stem, "summary", max_length=96)
        cache_dir = zotero_link_cache_root() / cache_dir_name
        ensure_dir(cache_dir.parent)
        copy_tree_contents(package_dir, cache_dir)
        packaged_md = cache_dir / f"{summary_stem}.md"
        if packaged_md.exists():
            return packaged_md.resolve()
        md_candidates = sorted(cache_dir.glob("*.md"))
        if md_candidates:
            return md_candidates[0].resolve()
    return Path(record["summary_md"]).resolve()


def display_authors_from_creators(creators: list[dict[str, str]] | None) -> list[str]:
    names: list[str] = []
    if not creators:
        return names
    for creator in creators:
        if not isinstance(creator, dict):
            continue
        name = first_nonempty(creator.get("name"))
        if name:
            names.append(name)
            continue
        family = first_nonempty(creator.get("family"), creator.get("lastName"))
        given = first_nonempty(creator.get("given"), creator.get("firstName"))
        combined = " ".join(part for part in (given, family) if part)
        if combined:
            names.append(combined)
    return names


def http_json(
    url: str,
    *,
    method: str = "GET",
    headers: dict[str, str] | None = None,
    json_body: Any | None = None,
    form_body: dict[str, Any] | None = None,
    raw_body: bytes | None = None,
    timeout: int = 120,
) -> tuple[int, dict[str, str], Any]:
    request_headers = dict(headers or {})
    data: bytes | None = None
    if json_body is not None:
        data = json.dumps(json_body, ensure_ascii=False).encode("utf-8")
        request_headers.setdefault("Content-Type", "application/json")
    elif form_body is not None:
        data = urllib.parse.urlencode(form_body).encode("utf-8")
        request_headers.setdefault("Content-Type", "application/x-www-form-urlencoded")
    elif raw_body is not None:
        data = raw_body

    request = urllib.request.Request(url, data=data, headers=request_headers, method=method)
    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            body = response.read()
            response_headers = {k: v for k, v in response.headers.items()}
            content_type = response.headers.get("Content-Type", "")
            if "application/json" in content_type:
                payload = json.loads(body.decode("utf-8"))
            else:
                payload = body
            return response.status, response_headers, payload
    except urllib.error.HTTPError as exc:
        body = exc.read().decode("utf-8", errors="replace")
        raise SyncError(f"HTTP {exc.code} for {url}: {body[:1200]}") from exc
    except urllib.error.URLError as exc:
        raise SyncError(f"Request failed for {url}: {exc}") from exc


def multipart_form_data(fields: dict[str, str], file_field: str, filename: str, data: bytes) -> tuple[str, bytes]:
    boundary = f"----CodexBoundary{uuid.uuid4().hex}"
    parts: list[bytes] = []
    for key, value in fields.items():
        parts.append(f"--{boundary}\r\n".encode("utf-8"))
        parts.append(
            f'Content-Disposition: form-data; name="{key}"\r\n\r\n{value}\r\n'.encode("utf-8")
        )
    mime_type = mimetypes.guess_type(filename)[0] or "application/octet-stream"
    parts.append(f"--{boundary}\r\n".encode("utf-8"))
    parts.append(
        (
            f'Content-Disposition: form-data; name="{file_field}"; filename="{filename}"\r\n'
            f"Content-Type: {mime_type}\r\n\r\n"
        ).encode("utf-8")
    )
    parts.append(data)
    parts.append(b"\r\n")
    parts.append(f"--{boundary}--\r\n".encode("utf-8"))
    return f"multipart/form-data; boundary={boundary}", b"".join(parts)


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def has_openai_api_key(config: Config) -> bool:
    return bool(os.environ.get(config.openai_api_key_env, "").strip())


def has_compatible_api_key(config: Config) -> bool:
    return bool(os.environ.get(config.compatible_api_key_env, "").strip())


def default_available_backend(config: Config) -> str:
    if config.openai_enabled and has_openai_api_key(config):
        return "openai"
    if config.compatible_enabled and has_compatible_api_key(config):
        return "openai_compatible"
    return "codex"


CODEX_CLI_EXECUTABLE_NAMES = ("codex", "codex.exe")


def normalize_codex_cli_command(command: str | None) -> str:
    return str(command or "").strip().strip('"')


def preferred_codex_platform_prefixes() -> tuple[str, ...]:
    if os.name == "nt":
        return ("windows-",)
    if sys.platform == "darwin":
        return ("darwin-", "macos-")
    return ("linux-",)


def codex_binary_platform_tag(path: Path) -> str | None:
    parts = [part.lower() for part in path.parts]
    try:
        bin_index = parts.index("bin")
    except ValueError:
        return None
    if bin_index + 1 >= len(parts):
        return None
    candidate = parts[bin_index + 1]
    if any(candidate.startswith(prefix) for prefix in ("windows-", "linux-", "darwin-", "macos-")):
        return candidate
    return None


def codex_path_is_platform_compatible(path: Path) -> bool:
    tag = codex_binary_platform_tag(path)
    if tag is None:
        return True
    return any(tag.startswith(prefix) for prefix in preferred_codex_platform_prefixes())


def codex_binary_name_matches(path: Path) -> bool:
    name = path.name.lower()
    if os.name == "nt":
        return name == "codex.exe"
    return name == "codex"


def find_installed_codex_cli() -> str | None:
    search_roots = (
        Path.home() / ".vscode" / "extensions",
        Path.home() / ".vscode-insiders" / "extensions",
    )
    candidates: list[Path] = []
    for root in search_roots:
        if not root.exists():
            continue
        for platform_prefix in preferred_codex_platform_prefixes():
            for candidate in root.glob(f"openai.chatgpt-*/bin/{platform_prefix}*/codex*"):
                try:
                    if (
                        candidate.is_file()
                        and codex_binary_name_matches(candidate)
                        and codex_path_is_platform_compatible(candidate)
                    ):
                        candidates.append(candidate.resolve())
                except OSError:
                    continue

    if not candidates:
        return None

    try:
        candidates.sort(key=lambda path: path.stat().st_mtime, reverse=True)
    except OSError:
        candidates.sort(reverse=True)
    return str(candidates[0])


def resolve_codex_cli_command(command: str | None) -> str | None:
    normalized = normalize_codex_cli_command(command)
    candidates: list[str] = []
    seen: set[str] = set()

    def add_candidate(value: str | None) -> None:
        resolved = normalize_codex_cli_command(value)
        key = resolved.lower()
        if not resolved or key in seen:
            return
        seen.add(key)
        candidates.append(resolved)

    add_candidate(normalized)
    command_name = Path(normalized).name.lower() if normalized else ""
    if not normalized or command_name in CODEX_CLI_EXECUTABLE_NAMES:
        for executable_name in CODEX_CLI_EXECUTABLE_NAMES:
            add_candidate(executable_name)

    for candidate in candidates:
        located = which(candidate)
        if located:
            located_path = Path(located)
            if codex_binary_name_matches(located_path) and codex_path_is_platform_compatible(located_path):
                return str(located_path.resolve())
            continue
        candidate_path = Path(candidate).expanduser()
        try:
            if (
                candidate_path.is_file()
                and codex_binary_name_matches(candidate_path)
                and codex_path_is_platform_compatible(candidate_path)
            ):
                return str(candidate_path.resolve())
        except OSError:
            continue

    return find_installed_codex_cli()


def codex_cli_available(config: Config) -> bool:
    if not config.codex_cli_enabled:
        return False
    return resolve_codex_cli_command(config.codex_cli_command) is not None


def resolve_path(base: Path, value: str | None, default: Path) -> Path:
    if not value:
        return default.resolve()
    candidate = Path(value)
    if not candidate.is_absolute():
        candidate = (base / candidate).resolve()
    return candidate.resolve()


def clone_config(config: Config, **updates: Any) -> Config:
    payload = dict(config.__dict__)
    payload.update(updates)
    return Config(**payload)


def apply_runtime_model_overrides(config: Config, args: argparse.Namespace) -> Config:
    backend = str(getattr(args, "summary_backend", "") or config.summary_backend_preference or "auto").strip().lower()
    if backend not in {"auto", "openai", "openai_compatible", "codex"}:
        backend = "auto"

    updated: dict[str, Any] = {"summary_backend_preference": backend}
    selected_backend = backend
    if selected_backend == "auto":
        selected_backend = default_available_backend(config)

    model_override = str(getattr(args, "summary_model", "") or "").strip()
    if model_override:
        if selected_backend == "codex":
            updated["codex_cli_model"] = model_override
        elif selected_backend == "openai_compatible":
            updated["compatible_model"] = model_override
        else:
            updated["openai_model"] = model_override

    reasoning_override = str(getattr(args, "summary_reasoning_effort", "") or "").strip().lower()
    if reasoning_override:
        if selected_backend == "codex":
            updated["codex_cli_reasoning_effort"] = reasoning_override
        elif selected_backend == "openai_compatible":
            pass
        else:
            updated["openai_reasoning_effort"] = reasoning_override

    updated["summary_user_request"] = str(getattr(args, "summary_user_request", "") or "").strip()

    return clone_config(config, **updated)


def load_config(path: Path) -> Config:
    config_path = path.expanduser().resolve()
    if not config_path.exists():
        raise FileNotFoundError(f"Config not found: {config_path}")
    raw = load_json(config_path, {})
    base_dir = config_path.parent
    workspace_root = resolve_path(base_dir, raw.get("workspace_root"), base_dir.parent)
    paper_dir = resolve_path(base_dir, raw.get("paper_dir"), workspace_root / "01-paper")
    summary_root = resolve_path(base_dir, raw.get("summary_root"), workspace_root)
    spec_path = resolve_path(
        base_dir,
        raw.get("spec_path"),
        workspace_root / "02-paper_summary_specs" / "default.md",
    )
    cache_root = resolve_path(base_dir, raw.get("cache_root"), base_dir / ".state")
    state_path = resolve_path(base_dir, raw.get("state_path"), base_dir / ".state" / "sync_state.json")
    export_path = resolve_path(base_dir, raw.get("zotero_export", {}).get("path"), base_dir / "exports" / "01-paper-sync.json")
    obsidian_vault_dir = resolve_path(
        base_dir,
        raw.get("obsidian", {}).get("vault_dir"),
        base_dir / "obsidian_vault",
    )
    zotero_local_raw = raw.get("zotero_local", {})
    zotero_local_db_value = zotero_local_raw.get("db_path")
    zotero_local_db_path: Path | None = None
    if zotero_local_db_value:
        zotero_local_db_path = resolve_path(base_dir, zotero_local_db_value, Path(zotero_local_db_value))
    collection_path = str(raw.get("zotero_export", {}).get("collection_path", "//01-paper-sync"))
    collection_parts = [part for part in collection_path.split("/") if part]
    inferred_collection_name = collection_parts[-1] if collection_parts else "01-paper-sync"

    tools_root = workspace_root / "03-tools"
    if not tools_root.exists():
        tools_root = workspace_root / "tools"
    tools_python = tools_root / "pdf_tools" / ".venv" / "Scripts" / "python.exe"
    if not tools_python.exists():
        tools_python = Path(sys.executable)
    extract_script = tools_root / "pdf_tools" / "extract_pdf.py"
    package_script = tools_root / "paper_workflow" / "package_summary.py"
    raw_codex_cli_command = raw.get("codex_cli", {}).get("command", "codex")
    resolved_codex_cli_command = resolve_codex_cli_command(str(raw_codex_cli_command))
    compatible_raw = raw.get("openai_compatible", {})
    compatible_base_url = str(compatible_raw.get("base_url", "https://api.deepseek.com")).rstrip("/")
    compatible_endpoint = str(
        compatible_raw.get("endpoint")
        or (compatible_base_url + "/v1/chat/completions")
    )

    return Config(
        config_path=config_path,
        workspace_root=workspace_root,
        paper_dir=paper_dir,
        summary_root=summary_root,
        spec_path=spec_path,
        state_path=state_path,
        cache_root=cache_root,
        export_path=export_path,
        poll_interval_seconds=int(raw.get("poll_interval_seconds", 30)),
        append_auto_materials=bool(raw.get("package_append_auto_materials", True)),
        openai_enabled=bool(raw.get("openai", {}).get("enabled", True)),
        openai_api_key_env=str(raw.get("openai", {}).get("api_key_env", "OPENAI_API_KEY")),
        openai_model=str(raw.get("openai", {}).get("model", "gpt-5-mini")),
        openai_reasoning_effort=raw.get("openai", {}).get("reasoning_effort", "medium"),
        openai_max_input_chars=int(raw.get("openai", {}).get("max_input_chars", 800000)),
        openai_endpoint=str(raw.get("openai", {}).get("endpoint", "https://api.openai.com/v1/responses")),
        compatible_enabled=bool(compatible_raw.get("enabled", True)),
        compatible_api_key_env=str(compatible_raw.get("api_key_env", "OPENAI_COMPATIBLE_API_KEY")),
        compatible_model=str(compatible_raw.get("model", "deepseek-chat")),
        compatible_base_url=compatible_base_url,
        compatible_endpoint=compatible_endpoint,
        compatible_temperature=float(compatible_raw.get("temperature", 0.2)),
        summary_backend_preference=str(raw.get("summary", {}).get("backend", "auto")).strip().lower() or "auto",
        codex_cli_enabled=bool(raw.get("codex_cli", {}).get("enabled", True)),
        codex_cli_command=resolved_codex_cli_command
        or normalize_codex_cli_command(str(raw_codex_cli_command))
        or "codex",
        codex_cli_model=raw.get("codex_cli", {}).get("model", None),
        codex_cli_reasoning_effort=raw.get("codex_cli", {}).get("reasoning_effort", None),
        zotero_api_enabled=bool(raw.get("zotero_api", {}).get("enabled", False)),
        zotero_api_key_env=str(raw.get("zotero_api", {}).get("api_key_env", "ZOTERO_API_KEY")),
        zotero_user_id=str(raw.get("zotero_api", {}).get("user_id", "")).strip(),
        zotero_library_type=str(raw.get("zotero_api", {}).get("library_type", "users")).strip() or "users",
        zotero_update_summary_note=bool(raw.get("zotero_api", {}).get("update_summary_note", True)),
        zotero_attach_summary_markdown=bool(raw.get("zotero_api", {}).get("attach_summary_markdown", True)),
        zotero_summary_attachment_mode=str(
            raw.get("zotero_api", {}).get("summary_attachment_mode", "linked_file")
        ).strip().lower()
        or "linked_file",
        obsidian_vault_dir=obsidian_vault_dir,
        obsidian_papers_subdir=str(raw.get("obsidian", {}).get("papers_subdir", "papers")),
        obsidian_tags_subdir=str(raw.get("obsidian", {}).get("tags_subdir", "tags")),
        obsidian_tag_prefix=str(raw.get("obsidian", {}).get("tag_prefix", "zotero")),
        bbt_jsonrpc_url=str(raw.get("zotero_export", {}).get("bbt_jsonrpc_url", "http://localhost:23119/better-bibtex/json-rpc")),
        bbt_collection_path=collection_path,
        bbt_translator=str(raw.get("zotero_export", {}).get("translator", "BetterBibTeX JSON")),
        bbt_display_options=dict(raw.get("zotero_export", {}).get("display_options", {})),
        bbt_replace=bool(raw.get("zotero_export", {}).get("replace", True)),
        zotero_local_enabled=bool(zotero_local_raw.get("enabled", True)),
        zotero_local_db_path=zotero_local_db_path,
        zotero_local_collection_name=str(
            zotero_local_raw.get("collection_name", inferred_collection_name)
        ).strip()
        or "01-paper-sync",
        zotero_local_fallback_to_all_pdf_items=bool(
            zotero_local_raw.get("fallback_to_all_pdf_items", True)
        ),
        tools_python=tools_python,
        extract_script=extract_script,
        package_script=package_script,
        summary_user_request="",
    )


def safe_unlink(path: Path) -> None:
    try:
        path.unlink()
    except FileNotFoundError:
        return


def discover_zotero_local_db(config: Config) -> Path | None:
    if config.zotero_local_db_path:
        return config.zotero_local_db_path if config.zotero_local_db_path.exists() else None

    candidates: list[Path] = []
    home = Path.home()
    candidates.append(home / "Zotero" / "zotero.sqlite")
    appdata = os.environ.get("APPDATA", "").strip()
    if appdata:
        candidates.extend(
            sorted((Path(appdata) / "Zotero" / "Zotero" / "Profiles").glob("*/zotero.sqlite"))
        )

    for candidate in candidates:
        if candidate.exists():
            return candidate.resolve()
    return None


def snapshot_zotero_local_db(config: Config, db_path: Path) -> Path:
    snapshot_root = config.cache_root / "zotero-local" / uuid.uuid4().hex
    ensure_dir(snapshot_root)
    snapshot_db = snapshot_root / "zotero.sqlite"
    shutil.copy2(db_path, snapshot_db)
    for suffix in ("-journal", "-wal", "-shm"):
        src = db_path.with_name(db_path.name + suffix)
        dst = snapshot_root / f"zotero.sqlite{suffix}"
        if src.exists():
            shutil.copy2(src, dst)
        else:
            safe_unlink(dst)
    return snapshot_db


def sqlite_placeholder_list(values: list[Any]) -> str:
    return ",".join("?" for _ in values)


def attachment_basename(path_value: str) -> str:
    text = path_value.replace("\\", "/").strip()
    if text.startswith("storage:"):
        return text.split("storage:", 1)[1].strip()
    if text.startswith("attachments:"):
        text = text.split("attachments:", 1)[1]
    return text.split("/")[-1].strip()


def load_export_records(path: Path) -> list[dict[str, Any]]:
    payload = load_json(path, [])
    items: list[Any]
    if isinstance(payload, list):
        items = payload
    elif isinstance(payload, dict):
        if isinstance(payload.get("items"), list):
            items = payload["items"]
        elif isinstance(payload.get("references"), list):
            items = payload["references"]
        else:
            values = list(payload.values())
            if values and all(isinstance(value, dict) for value in values):
                items = values
            else:
                items = [payload]
    else:
        items = []

    normalized: list[dict[str, Any]] = []
    for item in items:
        if not isinstance(item, dict):
            continue
        normalized_item = normalize_export_item(item)
        if normalized_item["title"]:
            normalized.append(normalized_item)
    return normalized


def chunked(values: list[Any], size: int = 500) -> list[list[Any]]:
    return [values[index : index + size] for index in range(0, len(values), size)]


def fetch_local_zotero_scope_item_ids(
    conn: sqlite3.Connection,
    *,
    collection_name: str,
    fallback_to_all_pdf_items: bool,
) -> tuple[list[int], str]:
    collection_rows = conn.execute(
        "select collectionID from collections where collectionName = ?",
        (collection_name,),
    ).fetchall()
    collection_ids = [int(row[0]) for row in collection_rows]
    if collection_ids:
        placeholders = sqlite_placeholder_list(collection_ids)
        scoped_rows = conn.execute(
            f"""
            select distinct i.itemID
            from collectionItems ci
            join items i on i.itemID = ci.itemID
            join itemTypes it on it.itemTypeID = i.itemTypeID
            where ci.collectionID in ({placeholders})
              and it.typeName not in ('attachment', 'note', 'annotation')
            """,
            collection_ids,
        ).fetchall()
        item_ids = [int(row[0]) for row in scoped_rows]
        if item_ids:
            return item_ids, f"collection:{collection_name}"
        log(f"[match] Local Zotero collection '{collection_name}' exists but is empty.")
    elif collection_name:
        log(f"[match] Local Zotero collection '{collection_name}' not found.")

    if not fallback_to_all_pdf_items:
        return [], "none"

    fallback_rows = conn.execute(
        """
        select distinct coalesce(ia.parentItemID, ia.itemID) as scopeItemID
        from itemAttachments ia
        join items i on i.itemID = coalesce(ia.parentItemID, ia.itemID)
        join itemTypes it on it.itemTypeID = i.itemTypeID
        where (
            lower(coalesce(ia.contentType, '')) = 'application/pdf'
            or lower(coalesce(ia.path, '')) like '%.pdf'
        )
          and it.typeName not in ('note', 'annotation')
        """
    ).fetchall()
    return [int(row[0]) for row in fallback_rows], "all-pdf-items"


def load_export_records_from_local_zotero(config: Config) -> list[dict[str, Any]]:
    if not config.zotero_local_enabled:
        return []

    db_path = discover_zotero_local_db(config)
    if db_path is None:
        log("[match] Local Zotero DB not found; skip runtime fallback.")
        return []

    snapshot_db = snapshot_zotero_local_db(config, db_path)
    conn = sqlite3.connect(snapshot_db)
    try:
        scope_item_ids, scope_label = fetch_local_zotero_scope_item_ids(
            conn,
            collection_name=config.zotero_local_collection_name,
            fallback_to_all_pdf_items=config.zotero_local_fallback_to_all_pdf_items,
        )
        if not scope_item_ids:
            log("[match] Local Zotero fallback found no candidate items.")
            return []

        item_rows: dict[int, dict[str, Any]] = {}
        field_map: dict[int, dict[str, str]] = {}
        creators_map: dict[int, list[dict[str, str]]] = {}
        tags_map: dict[int, list[str]] = {}
        attachments_map: dict[int, list[dict[str, str]]] = {}

        field_names = [
            "title",
            "DOI",
            "url",
            "archiveLocation",
            "publicationTitle",
            "journalAbbreviation",
            "archive",
            "publisher",
            "date",
            "citationKey",
        ]
        field_placeholders = sqlite_placeholder_list(field_names)

        for item_chunk in chunked(scope_item_ids):
            placeholders = sqlite_placeholder_list(item_chunk)
            for row in conn.execute(
                f"""
                select i.itemID, i.key, i.libraryID, it.typeName
                from items i
                join itemTypes it on it.itemTypeID = i.itemTypeID
                where i.itemID in ({placeholders})
                """,
                item_chunk,
            ).fetchall():
                item_rows[int(row[0])] = {
                    "item_id": int(row[0]),
                    "item_key": str(row[1]),
                    "library_id": row[2],
                    "item_type": str(row[3]),
                }

            field_params: list[Any] = list(item_chunk) + field_names
            for row in conn.execute(
                f"""
                select itemData.itemID, fields.fieldName, itemDataValues.value
                from itemData
                join fields on fields.fieldID = itemData.fieldID
                join itemDataValues on itemDataValues.valueID = itemData.valueID
                where itemData.itemID in ({placeholders})
                  and fields.fieldName in ({field_placeholders})
                """,
                field_params,
            ).fetchall():
                item_id = int(row[0])
                field_map.setdefault(item_id, {})[str(row[1])] = first_nonempty(row[2])

            for row in conn.execute(
                f"""
                select ic.itemID, c.firstName, c.lastName, c.fieldMode
                from itemCreators ic
                join creators c on c.creatorID = ic.creatorID
                where ic.itemID in ({placeholders})
                order by ic.itemID, ic.orderIndex
                """,
                item_chunk,
            ).fetchall():
                item_id = int(row[0])
                first_name = first_nonempty(row[1])
                last_name = first_nonempty(row[2])
                field_mode = int(row[3] or 0)
                if field_mode == 1:
                    creators_map.setdefault(item_id, []).append({"name": last_name})
                else:
                    creators_map.setdefault(item_id, []).append(
                        {"firstName": first_name, "lastName": last_name}
                    )

            for row in conn.execute(
                f"""
                select itemTags.itemID, tags.name
                from itemTags
                join tags on tags.tagID = itemTags.tagID
                where itemTags.itemID in ({placeholders})
                order by itemTags.itemID, tags.name
                """,
                item_chunk,
            ).fetchall():
                item_id = int(row[0])
                tag = first_nonempty(row[1])
                if tag:
                    tags_map.setdefault(item_id, []).append(tag)

            for row in conn.execute(
                f"""
                select coalesce(parentItemID, itemID) as ownerItemID, path
                from itemAttachments
                where coalesce(parentItemID, itemID) in ({placeholders})
                  and (
                    lower(coalesce(contentType, '')) = 'application/pdf'
                    or lower(coalesce(path, '')) like '%.pdf'
                  )
                order by ownerItemID, itemID
                """,
                item_chunk,
            ).fetchall():
                item_id = int(row[0])
                path_value = first_nonempty(row[1])
                if not path_value:
                    continue
                basename = attachment_basename(path_value)
                if not basename:
                    continue
                attachments_map.setdefault(item_id, []).append(
                    {"path": path_value, "basename": basename}
                )

        normalized: list[dict[str, Any]] = []
        for item_id, item in item_rows.items():
            fields = field_map.get(item_id, {})
            attachments = attachments_map.get(item_id, [])
            title = first_nonempty(fields.get("title"))
            if not title and item.get("item_type") == "attachment" and attachments:
                title = attachments[0]["basename"]
            raw_item = {
                "title": title,
                "DOI": first_nonempty(fields.get("DOI")),
                "url": first_nonempty(fields.get("url")),
                "archiveLocation": first_nonempty(fields.get("archiveLocation")),
                "publicationTitle": first_nonempty(fields.get("publicationTitle")),
                "journalAbbreviation": first_nonempty(fields.get("journalAbbreviation")),
                "archive": first_nonempty(fields.get("archive")),
                "publisher": first_nonempty(fields.get("publisher")),
                "date": first_nonempty(fields.get("date")),
                "citationKey": first_nonempty(fields.get("citationKey")),
                "key": item["item_key"],
                "itemKey": item["item_key"],
                "libraryID": item["library_id"],
                "creators": creators_map.get(item_id, []),
                "tags": sorted(dict.fromkeys(tags_map.get(item_id, []))),
                "attachments": attachments,
            }
            normalized_item = normalize_export_item(raw_item)
            if normalized_item["title"]:
                normalized.append(normalized_item)

        log(f"[match] Loaded {len(normalized)} Zotero records from local DB ({scope_label}).")
        return normalized
    finally:
        conn.close()


def normalize_creators(raw_creators: Any) -> list[dict[str, str]]:
    creators: list[dict[str, str]] = []
    if not isinstance(raw_creators, list):
        return creators
    for creator in raw_creators:
        if not isinstance(creator, dict):
            continue
        family = first_nonempty(creator.get("family"), creator.get("lastName"))
        given = first_nonempty(creator.get("given"), creator.get("firstName"))
        name = first_nonempty(creator.get("name"))
        if family or given:
            creators.append({"family": family, "given": given, "name": f"{given} {family}".strip()})
        elif name:
            creators.append({"family": name.split()[-1], "given": " ".join(name.split()[:-1]), "name": name})
    return creators


def normalize_tags(raw_tags: Any) -> list[str]:
    tags: list[str] = []
    if isinstance(raw_tags, str):
        parts = raw_tags.split(";") if ";" in raw_tags else [raw_tags]
        tags.extend(part.strip() for part in parts if part.strip())
    elif isinstance(raw_tags, list):
        for tag in raw_tags:
            if isinstance(tag, str):
                stripped = tag.strip()
                if stripped:
                    tags.append(stripped)
            elif isinstance(tag, dict):
                stripped = first_nonempty(tag.get("tag"), tag.get("name"))
                if stripped:
                    tags.append(stripped)
    return sorted(dict.fromkeys(tags))


def normalize_attachments(raw_attachments: Any) -> list[dict[str, str]]:
    attachments: list[dict[str, str]] = []
    if not isinstance(raw_attachments, list):
        return attachments
    for attachment in raw_attachments:
        if isinstance(attachment, str):
            attachments.append({"path": attachment, "basename": attachment_basename(attachment)})
            continue
        if not isinstance(attachment, dict):
            continue
        path = first_nonempty(
            attachment.get("path"),
            attachment.get("localPath"),
            attachment.get("filePath"),
            attachment.get("filename"),
        )
        basename = first_nonempty(
            attachment.get("basename"),
            attachment_basename(path) if path else "",
            attachment.get("title"),
            attachment.get("name"),
        )
        attachments.append(
            {
                "path": path,
                "basename": basename,
                "title": first_nonempty(attachment.get("title"), attachment.get("name")),
            }
        )
    return attachments


def extract_year_from_value(value: Any) -> str:
    if isinstance(value, dict):
        date_parts = value.get("date-parts")
        if isinstance(date_parts, list) and date_parts and isinstance(date_parts[0], list) and date_parts[0]:
            first = date_parts[0][0]
            if isinstance(first, int):
                return str(first)
            if isinstance(first, str) and YEAR_RE.search(first):
                return YEAR_RE.search(first).group(0)
    if isinstance(value, str):
        match = YEAR_RE.search(value)
        if match:
            return match.group(0)
    return ""


def normalize_export_item(raw: dict[str, Any]) -> dict[str, Any]:
    data = raw.get("data") if isinstance(raw.get("data"), dict) else raw
    title = first_nonempty(data.get("title"), raw.get("title"))
    creators = normalize_creators(
        data.get("creators") or raw.get("creators") or data.get("author") or raw.get("author")
    )
    tags = normalize_tags(data.get("tags") or raw.get("tags") or data.get("keyword") or raw.get("keyword"))
    attachments = normalize_attachments(data.get("attachments") or raw.get("attachments"))
    doi = first_nonempty(data.get("DOI"), raw.get("DOI"), data.get("doi"), raw.get("doi"))
    url = first_nonempty(data.get("url"), data.get("URL"), raw.get("url"), raw.get("URL"))
    arxiv = ""
    arxiv_match = ARXIV_RE.search(
        " ".join([url, title, first_nonempty(data.get("archiveLocation"), raw.get("archiveLocation"))])
    )
    if arxiv_match:
        arxiv = arxiv_match.group(1).lower()

    source = first_nonempty(
        data.get("publicationTitle"),
        data.get("container-title"),
        raw.get("container-title"),
        data.get("journalAbbreviation"),
        data.get("archive"),
        data.get("publisher"),
        raw.get("publicationTitle"),
    )
    year = first_nonempty(
        extract_year_from_value(data.get("issued")),
        extract_year_from_value(raw.get("issued")),
        extract_year_from_value(data.get("date")),
        extract_year_from_value(raw.get("date")),
        extract_year_from_value(data.get("year")),
        extract_year_from_value(raw.get("year")),
    )
    citation_key = first_nonempty(data.get("citationKey"), raw.get("citationKey"), raw.get("citation-key"))
    item_key = first_nonempty(data.get("key"), raw.get("key"), raw.get("itemKey"))
    library_id = data.get("libraryID") or raw.get("libraryID")

    return {
        "title": title,
        "title_norm": normalize_title(title),
        "creators": creators,
        "tags": tags,
        "attachments": attachments,
        "doi": normalize_identifier(doi),
        "url": url,
        "arxiv": arxiv,
        "source": source,
        "year": year,
        "citation_key": citation_key,
        "item_key": item_key,
        "library_id": library_id,
        "raw": raw,
    }


def build_export_index(records: list[dict[str, Any]]) -> dict[str, dict[str, list[dict[str, Any]]]]:
    attachment_map: dict[str, list[dict[str, Any]]] = {}
    doi_map: dict[str, list[dict[str, Any]]] = {}
    arxiv_map: dict[str, list[dict[str, Any]]] = {}
    title_records: list[dict[str, Any]] = []

    for record in records:
        title_records.append(record)
        if record["doi"]:
            doi_map.setdefault(record["doi"], []).append(record)
        if record["arxiv"]:
            arxiv_map.setdefault(record["arxiv"], []).append(record)
        for attachment in record["attachments"]:
            basename = first_nonempty(attachment.get("basename"))
            if basename:
                attachment_map.setdefault(basename.lower(), []).append(record)

    return {
        "attachment": attachment_map,
        "doi": doi_map,
        "arxiv": arxiv_map,
        "title": {"_": title_records},
    }


def run_subprocess(command: list[str], *, cwd: Path) -> str:
    result = subprocess.run(
        command,
        cwd=str(cwd),
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    if result.returncode != 0:
        message = result.stderr.strip() or result.stdout.strip() or "Unknown subprocess error"
        raise SyncError(message)
    return result.stdout


def _extended_path(p: Path) -> str:
    """Return a \\\\?\\  extended-length path on Windows to bypass MAX_PATH."""
    s = str(p)
    if s.startswith("\\\\?\\"):
        return s
    if s.startswith("\\\\"):
        return "\\\\?\\UNC\\" + s[2:]
    return "\\\\?\\" + s


def copy_tree_contents(src: Path, dst: Path) -> None:
    if os.name == "nt":
        # Use extended-length paths to handle long UNC paths (>260 chars)
        ext_dst = Path(_extended_path(dst.resolve()))
        ext_dst.mkdir(parents=True, exist_ok=True)
        command = [
            "robocopy",
            str(src),
            str(dst),
            "/MIR",
            "/R:1",
            "/W:1",
            "/NFL",
            "/NDL",
            "/NJH",
            "/NJS",
            "/NP",
        ]
        result = subprocess.run(
            command,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
        )
        if result.returncode >= 8:
            # Fallback to shutil with extended-length paths for long UNC paths
            log("[copy] robocopy failed, falling back to shutil.copytree")
            ext_src = _extended_path(src.resolve())
            ext_dst = _extended_path(dst.resolve())
            if dst.exists():
                shutil.rmtree(ext_dst)
            shutil.copytree(ext_src, ext_dst)
        return

    if dst.exists():
        shutil.rmtree(dst)
    shutil.copytree(src, dst)


def remove_tree_robust(path: Path) -> None:
    if not path.exists():
        return
    last_error: Exception | None = None
    target = _extended_path(path.resolve()) if os.name == "nt" else str(path)
    for _ in range(5):
        try:
            shutil.rmtree(target, ignore_errors=False)
            last_error = None
            break
        except Exception as exc:
            last_error = exc
            time.sleep(0.3)
    if path.exists() and last_error is not None:
        raise last_error


def extraction_dir_for_pdf(config: Config, pdf_path: Path) -> Path:
    digest = sha256_text(str(pdf_path.resolve()))[:10]
    return config.cache_root / "extracted" / f"{ascii_slug(pdf_path.stem)}-{digest}"


def extraction_cache_is_current(outdir: Path) -> bool:
    metadata_path = outdir / "metadata.json"
    document_md_path = outdir / "document.md"
    document_txt_path = outdir / "document.txt"
    if not metadata_path.is_file() or not document_md_path.is_file() or not document_txt_path.is_file():
        return False
    try:
        metadata = load_json(metadata_path, {})
    except Exception:
        return False
    version = metadata.get("extractor_version", 0)
    try:
        return int(version) >= EXTRACTOR_MIN_VERSION
    except (TypeError, ValueError):
        return False


def extract_pdf(config: Config, pdf_path: Path, *, force: bool) -> Path:
    outdir = extraction_dir_for_pdf(config, pdf_path)
    if outdir.exists() and not force and extraction_cache_is_current(outdir):
        return outdir
    ensure_dir(outdir.parent)
    command = [
        str(config.tools_python),
        str(config.extract_script),
        str(pdf_path),
        "--outdir",
        str(outdir),
        "--clean",
    ]
    log(f"[extract] {pdf_path.name}")
    run_subprocess(command, cwd=config.workspace_root)
    return outdir


def load_extraction_bundle(outdir: Path) -> dict[str, Any]:
    metadata = load_json(outdir / "metadata.json", {})
    document_md = read_text(outdir / "document.md") if (outdir / "document.md").exists() else ""
    document_txt = read_text(outdir / "document.txt") if (outdir / "document.txt").exists() else ""
    return {
        "metadata": metadata,
        "document_md": document_md,
        "document_txt": document_txt,
        "outdir": outdir,
    }


def infer_title_from_text(metadata: dict[str, Any], text: str) -> str:
    meta_title = first_nonempty(metadata.get("metadata", {}).get("title"), metadata.get("title"))
    if meta_title and meta_title.lower() not in {"extracted pdf", "untitled"}:
        return meta_title

    lines = []
    for raw_line in text.splitlines():
        stripped = raw_line.strip()
        if not stripped or stripped.startswith("====="):
            continue
        if stripped.lower().startswith("page "):
            continue
        lines.append(stripped)
        if len(lines) >= 15:
            break

    candidates = [
        line
        for line in lines[:8]
        if 12 <= len(line) <= 220 and not re.fullmatch(r"[A-Z0-9 .,:;()/-]+", line)
    ]
    if candidates:
        return max(candidates, key=len)
    return lines[0] if lines else ""


def infer_pdf_metadata(pdf_path: Path, bundle: dict[str, Any]) -> dict[str, Any]:
    metadata = bundle["metadata"]
    text = bundle["document_txt"]
    title = infer_title_from_text(metadata, text)
    doi_match = DOI_RE.search(text[:20000])
    arxiv_match = ARXIV_RE.search(" ".join([pdf_path.name, text[:20000]]))
    source = first_nonempty(
        metadata.get("metadata", {}).get("subject"),
        metadata.get("metadata", {}).get("keywords"),
    )
    return {
        "title": title,
        "title_norm": normalize_title(title),
        "doi": normalize_identifier(doi_match.group(0)) if doi_match else "",
        "arxiv": arxiv_match.group(1).lower() if arxiv_match else "",
        "source": source,
        "filename": pdf_path.name,
    }


def match_export_record(
    pdf_path: Path,
    inferred: dict[str, Any],
    export_index: dict[str, dict[str, list[dict[str, Any]]]],
) -> tuple[dict[str, Any] | None, float, str]:
    basename = pdf_path.name.lower()
    attachment_hits = export_index["attachment"].get(basename, [])
    if len(attachment_hits) == 1:
        return attachment_hits[0], 1.0, "attachment-basename"

    if inferred["doi"]:
        doi_hits = export_index["doi"].get(inferred["doi"], [])
        if len(doi_hits) == 1:
            return doi_hits[0], 0.99, "doi"

    if inferred["arxiv"]:
        arxiv_hits = export_index["arxiv"].get(inferred["arxiv"], [])
        if len(arxiv_hits) == 1:
            return arxiv_hits[0], 0.98, "arxiv"

    title = inferred["title"]
    title_norm = inferred["title_norm"]
    best_record: dict[str, Any] | None = None
    best_score = 0.0
    for candidate in export_index["title"]["_"]:
        score = difflib.SequenceMatcher(None, title_norm, candidate["title_norm"]).ratio()
        if score > best_score:
            best_score = score
            best_record = candidate
    if best_record and best_score >= 0.82:
        return best_record, best_score, "title-similarity"

    if attachment_hits:
        return attachment_hits[0], 0.9, "attachment-ambiguous-first"
    return None, 0.0, "unmatched"


def author_fragment(record: dict[str, Any] | None, pdf_path: Path) -> str:
    creators = record.get("creators") if record else []
    if creators:
        family = first_nonempty(creators[0].get("family"), creators[0].get("name"))
        fragment = filename_fragment(family, max_words=2, max_length=30)
        if fragment:
            return fragment
    return filename_fragment(pdf_path.stem, max_words=3, max_length=30) or "Paper"


def derive_summary_stem(pdf_path: Path, inferred: dict[str, Any], matched: dict[str, Any] | None) -> str:
    title = matched["title"] if matched else inferred["title"]
    year = matched["year"] if matched and matched["year"] else first_nonempty(
        YEAR_RE.search(pdf_path.stem).group(0) if YEAR_RE.search(pdf_path.stem) else "",
        str(time.localtime().tm_year),
    )
    source = matched["source"] if matched else inferred["source"]
    author = author_fragment(matched, pdf_path)
    title_part = filename_fragment(title or pdf_path.stem, max_words=8, max_length=48) or ascii_slug(pdf_path.stem)[:48]
    source_part = filename_fragment(source, max_words=4, max_length=18)
    parts = [author[:24], year[:4], title_part]
    if source_part:
        parts.append(source_part)
    stem = "-".join(part for part in parts if part)
    if len(stem) > 96:
        stem = stem[:96].rstrip("-")
    return stem


def build_summary_prompt(
    config: Config,
    pdf_path: Path,
    inferred: dict[str, Any],
    matched: dict[str, Any] | None,
    bundle: dict[str, Any],
    spec_text: str,
) -> list[dict[str, Any]]:
    metadata_payload = {
        "pdf_path": str(pdf_path.resolve()),
        "inferred_metadata": inferred,
        "matched_zotero_metadata": {
            "title": matched.get("title"),
            "year": matched.get("year"),
            "source": matched.get("source"),
            "creators": matched.get("creators"),
            "tags": matched.get("tags"),
            "doi": matched.get("doi"),
            "arxiv": matched.get("arxiv"),
            "citation_key": matched.get("citation_key"),
            "item_key": matched.get("item_key"),
        }
        if matched
        else None,
        "extraction_dir": str(bundle["outdir"]),
        "extraction_metadata": bundle["metadata"],
    }

    extracted_md = bundle["document_md"]
    truncated = False
    if len(extracted_md) > config.openai_max_input_chars:
        keep_head = int(config.openai_max_input_chars * 0.6)
        keep_tail = config.openai_max_input_chars - keep_head
        extracted_md = (
            extracted_md[:keep_head]
            + "\n\n[... 中间部分因输入长度限制被截断，以下继续保留文末内容 ...]\n\n"
            + extracted_md[-keep_tail:]
        )
        truncated = True

    system_prompt = (
        "你要为本地论文管理工作流生成中文 Markdown 总结。"
        "输出只能是最终 Markdown，不要加解释，不要加代码围栏。"
        "必须严格遵守给定规范，尤其是："
        "行内公式必须使用 $...$，块公式必须使用 $$...$$，"
        "“摘要翻译”必须紧跟在“一句话结论”之后，"
        "Method/Methods/附录/补充材料中的关键方法学内容必须纳入详细分析。"
        "如果原文没有直接回答某个点，可以明确写“原文未直接说明”，"
        "必要时再以“基于常识补充/推测”标记你的推断。"
    )

    user_prompt = (
        "请按照下面的最新默认规范，为给定 PDF 生成一份完整中文总结 Markdown。\n\n"
        f"最新规范全文：\n{spec_text}\n\n"
        f"本次用户补充要求和关键问题：\n{config.summary_user_request or '（无）'}\n\n"
        f"辅助元数据：\n{json.dumps(metadata_payload, ensure_ascii=False, indent=2)}\n\n"
        "写作要求补充：\n"
        "- 如果能从提取 Markdown 里识别到图片/图注/公式图片路径，可以在相关章节直接引用这些本地相对路径。\n"
        "- 如果不能可靠定位图片，就先保证文字总结和公式说明正确，打包阶段会补自动提取附录。\n"
        "- 不要省略 Methods、误差分析、资源估算、实验细节、补充材料中的关键技术信息。\n"
        "- 如果文中是理论推导，重点讲清假设、公式含义、推导逻辑和结论成立条件。\n"
        "- 不要把数学内容写成反引号代码样式。\n"
        f"- 本次输入是否截断：{'是' if truncated else '否'}。\n\n"
        f"以下是 PDF 自动提取得到的 Markdown 全文（或截断版本）：\n{extracted_md}\n"
    )

    return [
        {"role": "system", "content": [{"type": "input_text", "text": system_prompt}]},
        {"role": "user", "content": [{"type": "input_text", "text": user_prompt}]},
    ]


def extract_response_text(response_payload: Any) -> str:
    if isinstance(response_payload, dict):
        output_text = response_payload.get("output_text")
        if isinstance(output_text, str) and output_text.strip():
            return output_text
        outputs = response_payload.get("output")
        if isinstance(outputs, list):
            chunks: list[str] = []
            for item in outputs:
                if not isinstance(item, dict):
                    continue
                for content in item.get("content", []):
                    if not isinstance(content, dict):
                        continue
                    if content.get("type") in {"output_text", "text"}:
                        text = content.get("text")
                        if isinstance(text, str):
                            chunks.append(text)
            if chunks:
                return "".join(chunks)
    raise SyncError("OpenAI response did not contain output text.")


def prompt_text_from_payload(input_payload: list[dict[str, Any]]) -> str:
    parts: list[str] = []
    for message in input_payload:
        role = message.get("role", "user")
        parts.append(f"[{role}]")
        for content in message.get("content", []):
            if isinstance(content, dict) and content.get("type") == "input_text":
                parts.append(str(content.get("text", "")))
        parts.append("")
    return "\n".join(parts).strip() + "\n"


def call_openai(config: Config, input_payload: list[dict[str, Any]]) -> str:
    api_key = os.environ.get(config.openai_api_key_env, "").strip()
    if not api_key:
        raise SyncError(
            f"Missing environment variable {config.openai_api_key_env}; cannot auto-generate summaries."
        )
    log(
        f"[model] backend=openai model={config.openai_model} "
        f"reasoning={config.openai_reasoning_effort or 'default'}"
    )
    request_payload: dict[str, Any] = {"model": config.openai_model, "input": input_payload}
    if config.openai_reasoning_effort:
        request_payload["reasoning"] = {"effort": config.openai_reasoning_effort}

    _, _, response_payload = http_json(
        config.openai_endpoint,
        method="POST",
        headers={"Authorization": f"Bearer {api_key}"},
        json_body=request_payload,
        timeout=600,
    )
    return strip_markdown_fences(extract_response_text(response_payload))


def input_payload_to_chat_messages(input_payload: list[dict[str, Any]]) -> list[dict[str, str]]:
    messages: list[dict[str, str]] = []
    for message in input_payload:
        role = str(message.get("role") or "user")
        if role not in {"system", "user", "assistant"}:
            role = "user"
        chunks: list[str] = []
        for content in message.get("content", []):
            if isinstance(content, dict):
                text = content.get("text")
                if isinstance(text, str):
                    chunks.append(text)
        messages.append({"role": role, "content": "\n".join(chunks)})
    return messages


def extract_chat_completion_text(response_payload: Any) -> str:
    if isinstance(response_payload, dict):
        choices = response_payload.get("choices")
        if isinstance(choices, list) and choices:
            first = choices[0]
            if isinstance(first, dict):
                message = first.get("message")
                if isinstance(message, dict) and isinstance(message.get("content"), str):
                    return message["content"]
                text = first.get("text")
                if isinstance(text, str):
                    return text
    raise SyncError("OpenAI-compatible response did not contain message content.")


def call_openai_compatible(config: Config, input_payload: list[dict[str, Any]]) -> str:
    api_key = os.environ.get(config.compatible_api_key_env, "").strip()
    if not api_key:
        raise SyncError(
            f"Missing environment variable {config.compatible_api_key_env}; cannot use OpenAI-compatible backend."
        )
    log(
        f"[model] backend=openai_compatible model={config.compatible_model} "
        f"endpoint={config.compatible_endpoint}"
    )
    request_payload = {
        "model": config.compatible_model,
        "messages": input_payload_to_chat_messages(input_payload),
        "temperature": config.compatible_temperature,
    }
    _, _, response_payload = http_json(
        config.compatible_endpoint,
        method="POST",
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        json_body=request_payload,
        timeout=600,
    )
    return strip_markdown_fences(extract_chat_completion_text(response_payload))


def call_codex_cli(config: Config, input_payload: list[dict[str, Any]]) -> str:
    resolved_command = resolve_codex_cli_command(config.codex_cli_command)
    if not resolved_command:
        raise SyncError("codex CLI executable could not be resolved from config, PATH, or installed VS Code extensions.")

    state_dir = config.config_path.parent / ".state"
    ensure_dir(state_dir)
    output_path = state_dir / f"codex-exec-{uuid.uuid4().hex}.txt"
    output_rel = os.path.relpath(output_path, config.workspace_root).replace("\\", "/")
    prompt_text = prompt_text_from_payload(input_payload)

    command = [
        resolved_command,
        "exec",
        "--skip-git-repo-check",
        "-C",
        ".",
        "--color",
        "never",
        "--output-last-message",
        output_rel,
        "-",
    ]
    if config.codex_cli_model:
        command.extend(["-m", str(config.codex_cli_model)])
    if config.codex_cli_reasoning_effort:
        command.extend(["-c", f'model_reasoning_effort="{config.codex_cli_reasoning_effort}"'])

    log(
        f"[model] backend=codex model={config.codex_cli_model or 'config-default'} "
        f"reasoning={config.codex_cli_reasoning_effort or 'config-default'}"
    )
    log(f"[model] codex_executable={resolved_command}")

    try:
        result = subprocess.run(
            command,
            cwd=str(config.workspace_root),
            input=prompt_text,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
        )
    except OSError as exc:
        raise SyncError(f"Failed to launch codex CLI '{resolved_command}': {exc}") from exc
    if result.returncode != 0:
        message = result.stderr.strip() or result.stdout.strip() or "codex exec failed"
        raise SyncError(message)
    if not output_path.exists():
        raise SyncError("codex exec finished without writing the output message file.")

    text = output_path.read_text(encoding="utf-8")
    output_path.unlink(missing_ok=True)
    return strip_markdown_fences(text)


def call_model(config: Config, input_payload: list[dict[str, Any]]) -> str:
    backend = str(config.summary_backend_preference or "auto").strip().lower()
    if backend == "openai":
        return call_openai(config, input_payload)
    if backend == "openai_compatible":
        return call_openai_compatible(config, input_payload)
    if backend == "codex":
        return call_codex_cli(config, input_payload)
    if has_openai_api_key(config):
        return call_openai(config, input_payload)
    if config.compatible_enabled and has_compatible_api_key(config):
        return call_openai_compatible(config, input_payload)
    if codex_cli_available(config):
        return call_codex_cli(config, input_payload)
    raise SyncError(
        f"No {config.openai_api_key_env} found and codex CLI is unavailable. "
        "Configure an API key or use the manual prompt queue."
    )


def repair_summary_markdown(
    config: Config,
    spec_text: str,
    broken_markdown: str,
    package_error: str,
) -> str:
    system_prompt = (
        "你是 Markdown 总结修复器。"
        "你只能修复结构和格式问题，不能删掉核心内容。"
        "输出只能是修复后的最终 Markdown。"
    )
    user_prompt = (
        "下面是一份论文总结 Markdown。打包校验失败了。"
        "请基于报错，把它修成满足规范的版本。\n\n"
        f"最新规范全文：\n{spec_text}\n\n"
        f"报错信息：\n{package_error}\n\n"
        f"待修复 Markdown：\n{broken_markdown}\n"
    )
    return call_model(
        config,
        [
            {"role": "system", "content": [{"type": "input_text", "text": system_prompt}]},
            {"role": "user", "content": [{"type": "input_text", "text": user_prompt}]},
        ],
    )


def package_summary(
    config: Config,
    summary_md: Path,
    pdf_path: Path | list[Path] | tuple[Path, ...],
    *,
    clean: bool,
    package_dir: Path | None = None,
    spec_path: Path | None = None,
) -> Path:
    package_dir = package_dir or (config.summary_root / summary_md.stem)
    legacy_nested_dir = package_dir / package_dir.name
    if legacy_nested_dir.exists():
        log(f"[package] removing legacy nested package dir: {legacy_nested_dir.name}")
        remove_tree_robust(legacy_nested_dir)
    if isinstance(pdf_path, (list, tuple)):
        pdf_paths = [Path(p) for p in pdf_path]
    else:
        pdf_paths = [Path(pdf_path)]
    command = [
        str(config.tools_python),
        str(config.package_script),
        str(summary_md),
        "--outdir",
        str(package_dir),
        "--spec",
        str(spec_path or config.spec_path),
        "--pdf",
        *[str(path) for path in pdf_paths],
    ]
    if clean:
        command.append("--clean")
    if config.append_auto_materials:
        command.append("--append-auto-materials")
    log(f"[package] {summary_md.name}")
    run_subprocess(command, cwd=config.workspace_root)
    return package_dir


def ensure_summary_and_package(
    config: Config,
    pdf_path: Path,
    inferred: dict[str, Any],
    matched: dict[str, Any] | None,
    bundle: dict[str, Any],
    spec_text: str,
    *,
    force_regenerate: bool,
    existing_record: dict[str, Any] | None,
) -> tuple[Path, Path]:
    if existing_record:
        existing_summary = resolve_existing_path(existing_record.get("summary_md"), kind="file")
        if existing_summary is not None and not force_regenerate:
            try:
                package_dir = package_summary(config, existing_summary, pdf_path, clean=True)
                return existing_summary, package_dir
            except Exception as exc:
                log(f"[warn] Existing summary packaging failed, will regenerate: {exc}")

    summary_stem = derive_summary_stem(pdf_path, inferred, matched)
    summary_md = config.summary_root / f"{summary_stem}.md"
    if summary_md.exists() and not force_regenerate:
        try:
            package_dir = package_summary(config, summary_md, pdf_path, clean=True)
            return summary_md, package_dir
        except Exception as exc:
            log(f"[warn] Existing summary reuse failed, will regenerate: {exc}")

    if not config.openai_enabled and not codex_cli_available(config):
        raise SyncError("No enabled model backend is available for auto-summary generation.")

    prompt = build_summary_prompt(config, pdf_path, inferred, matched, bundle, spec_text)
    log(f"[summarize] {pdf_path.name} -> {summary_md.name}")
    summary_text = call_model(config, prompt)
    summary_md.write_text(summary_text, encoding="utf-8")

    try:
        package_dir = package_summary(config, summary_md, pdf_path, clean=True)
        return summary_md, package_dir
    except Exception as exc:
        log(f"[repair] {summary_md.name}")
        repaired = repair_summary_markdown(config, spec_text, summary_text, str(exc))
        summary_md.write_text(repaired, encoding="utf-8")
        package_dir = package_summary(config, summary_md, pdf_path, clean=True)
        return summary_md, package_dir


def build_obsidian_frontmatter(record: dict[str, Any]) -> str:
    authors = record.get("authors", [])
    if not isinstance(authors, list):
        authors = []
    lines = [
        "---",
        "codex_sync: true",
        f"title: {yaml_quote(record.get('title') or record.get('summary_stem') or '')}",
        f"summary_stem: {yaml_quote(record.get('summary_stem', ''))}",
        f"summary_available: {'true' if bool(record.get('summary_md')) and Path(record.get('summary_md', '')).exists() else 'false'}",
        f"pdf_path: {yaml_quote(record.get('pdf_path', ''))}",
        f"package_dir: {yaml_quote(record.get('package_dir', ''))}",
        f"summary_md: {yaml_quote(record.get('summary_md', ''))}",
        f"generated_at: {yaml_quote(record.get('generated_at', ''))}",
        f"date_added: {yaml_quote(record.get('date_added', ''))}",
        f"date_modified: {yaml_quote(record.get('date_modified', ''))}",
        f"year: {yaml_quote(record.get('year', ''))}",
        f"source: {yaml_quote(record.get('source', ''))}",
        f"doi: {yaml_quote(record.get('doi', ''))}",
        f"arxiv: {yaml_quote(record.get('arxiv', ''))}",
        f"zotero_item_key: {yaml_quote(record.get('zotero_item_key', ''))}",
        f"citation_key: {yaml_quote(record.get('citation_key', ''))}",
        "authors:",
    ]
    for author in authors:
        lines.append(f"  - {yaml_quote(str(author))}")
    lines.extend(
        [
        "tags:",
        ]
    )
    for tag in record.get("obsidian_tags", []):
        lines.append(f"  - {yaml_quote(tag)}")
    lines.append("ai_tags:")
    for tag in record.get("ai_tags", []):
        lines.append(f"  - {yaml_quote(tag)}")
    lines.append("zotero_tags_original:")
    for tag in record.get("tags", []):
        lines.append(f"  - {yaml_quote(tag)}")
    lines.append("---")
    return "\n".join(lines) + "\n\n"


def strip_frontmatter_if_generated(text: str) -> str:
    if not text.startswith("---\n"):
        return text
    end = text.find("\n---\n", 4)
    if end == -1:
        return text
    frontmatter = text[: end + 5]
    if "codex_sync: true" not in frontmatter:
        return text
    return text[end + 5 :].lstrip("\n")


def copy_primary_pdf_alias(record: dict[str, Any], target_dir: Path) -> str:
    pdf_path_value = first_nonempty(record.get("pdf_path"))
    if not pdf_path_value:
        return ""
    source_pdf = Path(pdf_path_value)
    if not source_pdf.exists():
        return ""
    suffix = source_pdf.suffix or ".pdf"
    alias_name = sanitize_windows_filename(
        f"{record.get('title') or source_pdf.stem}{suffix}",
        f"{record.get('summary_stem') or source_pdf.stem}{suffix}",
        max_length=96,
    )
    target_pdf = target_dir / alias_name
    if source_pdf.resolve() == target_pdf.resolve():
        return str(target_pdf)
    try:
        if target_pdf.exists():
            source_stat = source_pdf.stat()
            target_stat = target_pdf.stat()
            if (
                int(source_stat.st_size) == int(target_stat.st_size)
                and int(source_stat.st_mtime_ns) == int(target_stat.st_mtime_ns)
            ):
                return str(target_pdf)
        shutil.copy2(source_pdf, target_pdf)
    except PermissionError:
        if target_pdf.exists():
            return str(target_pdf)
        return ""
    except OSError:
        if target_pdf.exists():
            return str(target_pdf)
        return ""
    return str(target_pdf)


def derive_obsidian_item_stem(record: dict[str, Any]) -> str:
    existing = first_nonempty(record.get("summary_stem"))
    if existing and len(existing) <= 72:
        return existing
    authors = record.get("authors", [])
    lead_author = ""
    if isinstance(authors, list) and authors:
        lead_author = filename_fragment(str(authors[0]), 3, 18)
    if not lead_author:
        lead_author = filename_fragment(record.get("title", ""), 3, 18)
    year = first_nonempty(record.get("year"), "n.d.")
    title_fragment = filename_fragment(record.get("title", ""), 8, 40)
    key = first_nonempty(record.get("zotero_item_key"), record.get("citation_key"), "item")[:8]
    stem = ascii_slug(f"{lead_author}-{year}-{title_fragment}-{key}")
    return stem[:72].strip("-") or f"item-{key}"


def build_obsidian_metadata_body(record: dict[str, Any], *, renamed_pdf_name: str) -> str:
    title = record.get("title") or record.get("summary_stem") or "Untitled"
    lines = [f"# {title}", ""]
    if not (record.get("summary_md") and Path(record.get("summary_md", "")).exists()):
        lines.append("> This is a metadata stub generated from Zotero. No AI summary markdown is attached yet.")
        lines.append("")
    lines.append("## Metadata")
    authors = record.get("authors", [])
    lines.append(f"- Authors: {', '.join(str(author) for author in authors) or 'N/A'}")
    lines.append(f"- Year: {record.get('year') or 'N/A'}")
    lines.append(f"- Source: {record.get('source') or 'N/A'}")
    lines.append(f"- DOI: {record.get('doi') or 'N/A'}")
    lines.append(f"- arXiv: {record.get('arxiv') or 'N/A'}")
    lines.append(f"- Zotero Item Key: {record.get('zotero_item_key') or 'N/A'}")
    lines.append(f"- Citation Key: {record.get('citation_key') or 'N/A'}")
    lines.append(f"- Date Added: {record.get('date_added') or 'N/A'}")
    lines.append(f"- Date Modified: {record.get('date_modified') or 'N/A'}")
    if renamed_pdf_name:
        lines.append(f"- PDF: [{renamed_pdf_name}]({urllib.parse.quote(renamed_pdf_name)})")
    elif record.get("pdf_path"):
        lines.append(f"- PDF Path: `{record.get('pdf_path')}`")
    summary_md = first_nonempty(record.get("summary_md"))
    if summary_md and Path(summary_md).exists():
        summary_name = Path(summary_md).name
        lines.append(f"- Summary Markdown: [{summary_name}]({urllib.parse.quote(summary_name)})")
    lines.append("")
    lines.append("## Tags")
    tags = record.get("tags", [])
    if tags:
        for tag in tags:
            lines.append(f"- {tag}")
    else:
        lines.append("- N/A")
    lines.append("")
    return "\n".join(lines).rstrip() + "\n"


# ── AI-tag extraction from summary Markdown ──────────────────────────────

_AI_TAGS_RE = re.compile(
    r"[-\*]\s*关键\s*tags\s*[：:]+\s*(.+)",
    re.IGNORECASE,
)


def extract_ai_tags_from_md(md_path: Path) -> list[str]:
    """Extract AI-generated tags from the '关键tags' field in 论文信息.

    Expected format in the Markdown::

        - 关键tags：tag1, tag2, tag3, tag4, tag5
        - 关键tags：tag1、tag2、tag3、tag4、tag5

    Returns a deduplicated, sorted list of stripped tag strings.
    """
    if not md_path.exists():
        return []
    try:
        text = read_text(md_path)
    except Exception:
        return []
    for line in text.splitlines():
        m = _AI_TAGS_RE.search(line)
        if m:
            raw = m.group(1).strip()
            # split on Chinese/English comma or semicolon
            parts = re.split(r"[,，;；、]+", raw)
            tags = [t.strip().strip("#") for t in parts if t.strip()]
            return sorted(dict.fromkeys(tags))
    return []


_ATTACHMENT_SUBDIRS = {"extracted", "images", "sources"}
_ATTACHMENTS_DIR_NAME = "_attachments"


def _resolve_legacy_asset_ref(target_dir: Path, ref: str) -> str:
    attachments_dir = target_dir / _ATTACHMENTS_DIR_NAME
    if not attachments_dir.exists():
        return ref

    cleaned = ref.strip().strip("<>").split("#", 1)[0]
    cleaned = urllib.parse.unquote(cleaned)
    if not cleaned.startswith("assets/"):
        return ref

    preferred_matches: list[Path] = []
    basename = Path(cleaned).name
    if basename:
        for candidate in attachments_dir.rglob(basename):
            try:
                relative = candidate.relative_to(attachments_dir)
            except ValueError:
                continue
            if "/assets/" not in str(relative).replace("\\", "/"):
                continue
            preferred_matches.append(candidate)

    if not preferred_matches:
        return ref

    chosen = sorted(preferred_matches, key=lambda value: len(str(value)))[0]
    relative_path = chosen.relative_to(target_dir).as_posix()
    return urllib.parse.quote(relative_path, safe="/")


def _rewrite_attachment_refs(text: str, target_dir: Path) -> str:
    """Rewrite relative references so they point inside _attachments/.

    Handles both Markdown image/link syntax  ``](path)``  and
    inline-code back-tick references  `` `sources/...` ``.
    """
    prefix = _ATTACHMENTS_DIR_NAME
    # ](images/...) ](extracted/...) ](sources/...)
    text = re.sub(
        r'\]\((?!_attachments/)(images/|extracted/|sources/)',
        lambda m: f"]({prefix}/{m.group(1)}",
        text,
    )
    text = re.sub(
        r'\]\((?!_attachments/)(assets/[^)]+)\)',
        lambda m: f"]({_resolve_legacy_asset_ref(target_dir, m.group(1))})",
        text,
    )
    # `sources/...`  (back-tick code references)
    text = re.sub(
        r'`(?!_attachments/)(sources/[^`]+)`',
        lambda m: f"`{prefix}/{m.group(1)}`",
        text,
    )
    return text


def _elevate_md_files(attachments_dir: Path, target_dir: Path) -> None:
    """Move top-level .md files from _attachments/ up to paper root.

    Only moves files directly under *attachments_dir*, not nested ones
    like ``extracted/slug/document.md``.
    Uses copy+remove instead of move for better NAS compatibility.
    Uses extended-length paths on Windows to bypass MAX_PATH.
    """
    if not attachments_dir.exists():
        return
    for md in list(attachments_dir.glob("*.md")):
        dest = target_dir / md.name
        src_s = _extended_path(md) if os.name == "nt" else str(md)
        dst_s = _extended_path(dest) if os.name == "nt" else str(dest)
        shutil.copy2(src_s, dst_s)
        try:
            Path(src_s).unlink()
        except OSError:
            pass  # source removal failure is non-critical


def _rewrite_all_md_refs(target_dir: Path) -> None:
    """Rewrite attachment refs in every .md directly under *target_dir*."""
    for md_file in target_dir.glob("*.md"):
        original = read_text(md_file)
        updated = _rewrite_attachment_refs(original, target_dir)
        if updated != original:
            md_file.write_text(updated, encoding="utf-8")


def sync_obsidian_package(config: Config, record: dict[str, Any], *, copy_package: bool) -> dict[str, Any]:
    package_dir = Path(record["package_dir"])
    summary_stem = record.get("obsidian_stem") or record["summary_stem"]
    source_summary_stem = str(record.get("summary_stem") or "")
    vault_dir = config.obsidian_vault_dir
    papers_root = vault_dir / config.obsidian_papers_subdir
    target_dir = papers_root / summary_stem
    target_md = target_dir / f"{summary_stem}.md"
    attachments_dir = target_dir / _ATTACHMENTS_DIR_NAME
    log(f"[obsidian] {summary_stem}")

    ensure_dir(papers_root)
    ensure_dir(vault_dir / config.obsidian_tags_subdir)

    if copy_package or not target_dir.exists():
        # Copy the whole package into _attachments/
        copy_tree_contents(package_dir, attachments_dir)
        # Move top-level .md files up to paper root
        _elevate_md_files(attachments_dir, target_dir)

    source_summary = Path(record.get("summary_md", ""))
    packaged_primary_md = target_dir / f"{source_summary_stem}.md" if source_summary_stem else None
    if copy_package and packaged_primary_md is not None and packaged_primary_md.exists():
        body = read_text(packaged_primary_md)
    elif target_md.exists():
        body = read_text(target_md)
    elif packaged_primary_md is not None and packaged_primary_md.exists():
        body = read_text(packaged_primary_md)
    elif source_summary.exists():
        body = read_text(source_summary)
    else:
        raise FileNotFoundError(str(target_md))
    body = strip_frontmatter_if_generated(body)
    body = _rewrite_attachment_refs(body, target_dir)

    # Extract AI-generated tags from summary MD content
    ai_tags = extract_ai_tags_from_md(target_md) if target_md.exists() else []
    if not ai_tags:
        source_summary = Path(record.get("summary_md", ""))
        if source_summary.exists():
            ai_tags = extract_ai_tags_from_md(source_summary)
    if ai_tags:
        record["ai_tags"] = ai_tags
        # Merge into obsidian_tags with ai/ prefix
        existing = set(record.get("obsidian_tags", []))
        for t in ai_tags:
            ai_obs_tag = f"ai/{sanitize_obsidian_tag(t)}"
            if ai_obs_tag not in existing:
                existing.add(ai_obs_tag)
        record["obsidian_tags"] = sorted(existing)

    frontmatter = build_obsidian_frontmatter(record)
    ensure_dir(target_md.parent)
    target_md.write_text(frontmatter + body, encoding="utf-8")
    skipped_md_paths = {target_md.resolve()}
    if packaged_primary_md is not None and packaged_primary_md.exists() and packaged_primary_md != target_md:
        packaged_primary_md.unlink(missing_ok=True)
        skipped_md_paths.add(packaged_primary_md.resolve())

    # Rewrite refs in any additional .md files (e.g. Shor-Algorithm-Explained.md)
    for other_md in target_dir.glob("*.md"):
        try:
            resolved_other_md = other_md.resolve()
        except OSError:
            continue
        if resolved_other_md in skipped_md_paths:
            continue
        try:
            txt = read_text(other_md)
        except OSError:
            continue
        rewritten = _rewrite_attachment_refs(txt, target_dir)
        if rewritten != txt:
            other_md.write_text(rewritten, encoding="utf-8")

    record["obsidian_pdf"] = copy_primary_pdf_alias(record, target_dir)
    record["obsidian_dir"] = str(target_dir)
    record["obsidian_md"] = str(target_md)
    return record


def sync_obsidian_metadata_record(config: Config, record: dict[str, Any]) -> dict[str, Any]:
    summary_stem = record.get("obsidian_stem") or derive_obsidian_item_stem(record)
    record["summary_stem"] = summary_stem
    vault_dir = config.obsidian_vault_dir
    papers_root = vault_dir / config.obsidian_papers_subdir
    target_dir = papers_root / summary_stem
    target_md = target_dir / f"{summary_stem}.md"
    ensure_dir(target_dir)
    renamed_pdf = copy_primary_pdf_alias(record, target_dir)
    renamed_pdf_name = Path(renamed_pdf).name if renamed_pdf else ""
    frontmatter = build_obsidian_frontmatter(record)
    body = build_obsidian_metadata_body(record, renamed_pdf_name=renamed_pdf_name)
    target_md.write_text(frontmatter + body, encoding="utf-8")
    record["obsidian_pdf"] = renamed_pdf
    record["obsidian_dir"] = str(target_dir)
    record["obsidian_md"] = str(target_md)
    return record


def rebuild_obsidian_indexes(config: Config, records: list[dict[str, Any]]) -> None:
    vault_dir = config.obsidian_vault_dir
    tags_root = vault_dir / config.obsidian_tags_subdir
    for _ in range(3):
        if tags_root.exists():
            if tags_root.is_dir():
                shutil.rmtree(tags_root, ignore_errors=True)
            else:
                tags_root.unlink(missing_ok=True)
        try:
            ensure_dir(tags_root)
            break
        except FileExistsError:
            time.sleep(0.2)
    else:
        ensure_dir(tags_root)

    # Deduplicate records by obsidian_md path to avoid repeated entries
    seen_md: set[str] = set()
    existing_records: list[dict[str, Any]] = []
    for record in records:
        md_path = record.get("obsidian_md", "")
        if not md_path or not Path(md_path).exists():
            continue
        norm = str(Path(md_path).resolve())
        if norm in seen_md:
            continue
        seen_md.add(norm)
        existing_records.append(record)

    log(f"[obsidian] rebuild indexes ({len(existing_records)} notes)")

    tag_map: dict[str, list[dict[str, Any]]] = {}
    for record in existing_records:
        for tag in record.get("obsidian_tags", []):
            tag_map.setdefault(tag, []).append(record)

    for tag, items in sorted(tag_map.items()):
        tag_file = tags_root / f"{sanitize_obsidian_tag(tag)}.md"
        ensure_dir(tag_file.parent)
        lines = [f"# {tag}", ""]
        for item in sorted(items, key=lambda value: value.get("title", "").lower()):
            obsidian_md = Path(item["obsidian_md"])
            rel = os.path.relpath(obsidian_md, tag_file.parent).replace("\\", "/")
            lines.append(f"- [{item.get('title') or item.get('summary_stem')}]({rel})")
            lines.append(f"- 关键词：{', '.join(item.get('tags', [])) or '无'}")
            lines.append("")
        tag_file.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")

    library_note = vault_dir / "00-Library.md"
    lines = ["# Paper Library", ""]
    for record in sorted(existing_records, key=lambda value: value.get("title", "").lower()):
        obsidian_md = Path(record["obsidian_md"])
        rel = os.path.relpath(obsidian_md, library_note.parent).replace("\\", "/")
        lines.append(f"- [{record.get('title') or record.get('summary_stem')}]({rel})")
        lines.append(f"- 标签：{', '.join(record.get('obsidian_tags', [])) or '无'}")
        lines.append("")
    library_note.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")

    # ── Generate 00-Dashboard.md with Dataview queries + static table ──
    rebuild_dashboard(config, existing_records)


def _escape_pipe(text: str) -> str:
    """Escape pipe characters for Markdown table cells."""
    return text.replace("|", "\\|").replace("\n", " ")


def _build_dataviewjs_block() -> str:
    """Return the dataviewjs code for the interactive paper table."""
    return r'''const pages = dv.pages('"papers"').where(p => p.codex_sync === true);
const container = dv.container;

// --- Collect unique tags ---
const allTags = new Set();
for (const p of pages) {
  for (const t of (p.zotero_tags_original || [])) {
    if (!String(t).startsWith("/")) allTags.add(String(t));
  }
  for (const t of (p.ai_tags || [])) { allTags.add(String(t)); }
}
const tagList = [...allTags].sort();

// --- Controls ---
const ctrl = container.createEl("div", {attr:{style:"display:flex;gap:8px;margin-bottom:10px;flex-wrap:wrap;align-items:center;"}});
const searchInput = ctrl.createEl("input", {
  type: "text", placeholder: "\uD83D\uDD0D 搜索标题/作者/关键词\u2026",
  attr:{style:"flex:1;min-width:200px;padding:6px 10px;border:1px solid var(--background-modifier-border);border-radius:6px;background:var(--background-primary);color:var(--text-normal);font-size:14px;"}
});
const tagSelect = ctrl.createEl("select", {
  attr:{style:"padding:6px 10px;border:1px solid var(--background-modifier-border);border-radius:6px;background:var(--background-primary);color:var(--text-normal);font-size:14px;"}
});
tagSelect.createEl("option", {text: "全部标签", attr:{value:""}});
for (const tag of tagList) tagSelect.createEl("option", {text: tag, attr:{value: tag}});

// --- Sort state ---
let sortCol = "date_added", sortAsc = false;

function getVal(p, col) {
  if (col === "first_author") return String((p.authors || [])[0] || "");
  if (col === "tags_display") return [...(p.zotero_tags_original || []).filter(t => !String(t).startsWith("/")), ...(p.ai_tags || [])].join(", ");
  return String(p[col] || "");
}

function render() {
  const old = container.querySelector(".dvjs-wrap");
  if (old) old.remove();
  const ft = (searchInput.value || "").toLowerCase();
  const selTag = tagSelect.value;
  let data = [...pages];
  if (ft) data = data.filter(p => [p.title, ...(p.authors||[]), ...(p.zotero_tags_original||[]), ...(p.ai_tags||[]), p.source||""].join(" ").toLowerCase().includes(ft));
  if (selTag) data = data.filter(p => (p.zotero_tags_original||[]).map(String).includes(selTag) || (p.ai_tags||[]).map(String).includes(selTag));
  data.sort((a, b) => { let c = getVal(a,sortCol).localeCompare(getVal(b,sortCol),undefined,{numeric:true}); return sortAsc ? c : -c; });

  const wrap = container.createEl("div", {cls:"dvjs-wrap", attr:{style:"overflow-x:auto;"}});
  const tbl = wrap.createEl("table", {attr:{style:"width:100%;border-collapse:collapse;font-size:14px;"}});

  const cols = [
    {l:"#",c:null,w:"30px"}, {l:"标题 (PDF)",c:"title"}, {l:"AI总结",c:null,w:"60px"},
    {l:"第一作者",c:"first_author"}, {l:"年份",c:"year",w:"50px"}, {l:"期刊",c:"source"},
    {l:"状态",c:"summary_available",w:"50px"}, {l:"关键词",c:"tags_display"}, {l:"添加日期",c:"date_added",w:"100px"}
  ];
  const hr = tbl.createEl("thead").createEl("tr");
  for (const h of cols) {
    const th = hr.createEl("th", {
      text: h.l + (h.c && sortCol===h.c ? (sortAsc?" ▲":" ▼") : ""),
      attr:{style:"padding:6px 8px;border-bottom:2px solid var(--background-modifier-border);text-align:left;white-space:nowrap;"
        + (h.c?"cursor:pointer;":"") + (h.w?"width:"+h.w+";":"")}
    });
    if (h.c) th.addEventListener("click", () => { if (sortCol===h.c) sortAsc=!sortAsc; else {sortCol=h.c;sortAsc=true;} render(); });
  }

  const tbody = tbl.createEl("tbody");
  data.forEach((p, i) => {
    const tr = tbody.createEl("tr", {attr:{style:"border-bottom:1px solid var(--background-modifier-border);"+(i%2?" background:var(--background-secondary);":"")}});
    tr.createEl("td", {text:String(i+1), attr:{style:"padding:4px 8px;text-align:center;"}});

    // Title → PDF link
    const tdT = tr.createEl("td", {attr:{style:"padding:4px 8px;"}});
    const pp = String(p.pdf_path || "");
    if (pp) { const u="file:///"+pp.replace(/\\/g,"/").replace(/ /g,"%20"); tdT.createEl("a",{text:p.title||"Untitled",attr:{href:u,class:"external-link",style:"color:var(--text-accent);text-decoration:none;"}}); }
    else tdT.setText(p.title||"Untitled");

    // AI Summary → internal link
    const tdS = tr.createEl("td", {attr:{style:"padding:4px 8px;text-align:center;"}});
    if (p.summary_available) tdS.createEl("a",{text:"\uD83D\uDCDD",cls:"internal-link",attr:{"data-href":p.file.path,href:p.file.path,style:"text-decoration:none;font-size:1.2em;"}});
    else tdS.setText("—");

    const aus = p.authors||[];
    tr.createEl("td",{text:aus.length?String(aus[0]):"N/A",attr:{style:"padding:4px 8px;"}});
    tr.createEl("td",{text:String(p.year||""),attr:{style:"padding:4px 8px;text-align:center;"}});
    tr.createEl("td",{text:String(p.source||""),attr:{style:"padding:4px 8px;"}});
    tr.createEl("td",{text:p.summary_available?"✅":"❌",attr:{style:"padding:4px 8px;text-align:center;"}});
    const tags = [...(p.zotero_tags_original||[]).filter(t=>!String(t).startsWith("/")), ...(p.ai_tags||[])];
    tr.createEl("td",{text:tags.slice(0,5).join(", "),attr:{style:"padding:4px 8px;font-size:0.9em;"}});
    tr.createEl("td",{text:String(p.date_added||"").substring(0,10),attr:{style:"padding:4px 8px;white-space:nowrap;"}});
  });

  wrap.createEl("div",{attr:{style:"padding:4px 8px;font-size:12px;color:var(--text-muted);"}}).setText("显示 "+data.length+" / "+pages.length+" 篇");
}
searchInput.addEventListener("input", render);
tagSelect.addEventListener("change", render);
render();'''


def _pdf_file_uri(pdf_path: str) -> str:
    """Convert a Windows path to a file:/// URI suitable for Obsidian links."""
    return "file:///" + pdf_path.replace("\\", "/").replace(" ", "%20")


def rebuild_dashboard(config: Config, records: list[dict[str, Any]]) -> None:
    vault_dir = config.obsidian_vault_dir
    dashboard = vault_dir / "00-Dashboard.md"
    sorted_records = sorted(records, key=lambda r: r.get("date_added", ""), reverse=True)

    # Collect all unique tags (excluding /unread and meta tags)
    all_tags: set[str] = set()
    for record in records:
        for tag in record.get("obsidian_tags", []):
            if tag not in ("zotero/unread",):
                all_tags.add(tag)

    lines: list[str] = []
    lines.append("# 📄 Paper Dashboard")
    lines.append("")
    lines.append(f"> 自动生成于 {time.strftime('%Y-%m-%d %H:%M')} · 共 {len(records)} 篇论文")
    lines.append("> 需安装 **Dataview** 插件并开启 DataviewJS 查看交互式表格；下方同时提供静态 Markdown 表格。")
    lines.append("")

    # ── DataviewJS interactive section ──
    lines.append("## 交互式检索（DataviewJS）")
    lines.append("")
    lines.append("```dataviewjs")
    lines.append(_build_dataviewjs_block())
    lines.append("```")
    lines.append("")

    # ── Dataview: pending summaries ──
    lines.append("### 待总结论文")
    lines.append("")
    lines.append("```dataview")
    lines.append("TABLE WITHOUT ID")
    lines.append('  link(file.link, title) AS "标题",')
    lines.append('  authors[0] AS "第一作者",')
    lines.append('  year AS "年份",')
    lines.append('  dateformat(date(date_added), "yyyy-MM-dd") AS "添加日期"')
    lines.append('FROM "papers"')
    lines.append("WHERE codex_sync = true AND summary_available = false")
    lines.append("SORT date_added DESC")
    lines.append("```")
    lines.append("")

    # ── Static Markdown table (always visible) ──
    lines.append("---")
    lines.append("")
    lines.append("## 静态总览表")
    lines.append("")
    lines.append("| # | 标题 (PDF) | AI总结 | 第一作者 | 年份 | 期刊 | 总结 | 关键词 | 添加日期 |")
    lines.append("|---|-----------|--------|----------|------|------|------|--------|----------|")

    for idx, record in enumerate(sorted_records, 1):
        title = record.get("title") or record.get("summary_stem") or "Untitled"

        # Title → PDF link
        pdf_path = record.get("pdf_path", "")
        if pdf_path and Path(pdf_path).exists():
            title_link = f"[{_escape_pipe(title)}]({_pdf_file_uri(pdf_path)})"
        else:
            title_link = _escape_pipe(title)

        # AI Summary → Obsidian markdown link
        obsidian_md = record.get("obsidian_md", "")
        if obsidian_md and Path(obsidian_md).exists():
            rel = os.path.relpath(Path(obsidian_md), dashboard.parent).replace("\\", "/")
            summary_link = f"[📝]({rel})"
        else:
            summary_link = "—"

        authors = record.get("authors", [])
        first_author = _escape_pipe(str(authors[0])) if authors else "N/A"

        year = record.get("year") or ""
        source = _escape_pipe(record.get("source") or "")
        has_summary = "✅" if record.get("summary_md") and Path(record.get("summary_md", "")).exists() else "❌"

        raw_tags = record.get("tags", [])
        display_tags = [t for t in raw_tags if not t.startswith("/")]
        tags_str = _escape_pipe(", ".join(display_tags[:3])) if display_tags else ""

        date_added = record.get("date_added", "")[:10]

        lines.append(f"| {idx} | {title_link} | {summary_link} | {first_author} | {year} | {source} | {has_summary} | {tags_str} | {date_added} |")

    lines.append("")

    # ── Statistics ──
    summarized = sum(1 for r in records if r.get("summary_md") and Path(r.get("summary_md", "")).exists())
    pending = len(records) - summarized
    lines.append("---")
    lines.append("")
    lines.append("## 统计")
    lines.append("")
    lines.append(f"- 总计：**{len(records)}** 篇")
    lines.append(f"- 已总结：**{summarized}** 篇")
    lines.append(f"- 待总结：**{pending}** 篇")
    lines.append(f"- 标签分类：**{len(all_tags)}** 个")
    lines.append("")

    dashboard.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    log(f"[obsidian] dashboard updated: {len(records)} papers")


def cleanup_stale_obsidian_folders(config: Config, records: list[dict[str, Any]]) -> None:
    """Remove obsidian_vault/papers/ folders not referenced by any current record."""
    papers_root = config.obsidian_vault_dir / config.obsidian_papers_subdir
    if not papers_root.exists():
        return

    expected: set[str] = set()
    for record in records:
        obsidian_dir = record.get("obsidian_dir", "")
        if obsidian_dir:
            expected.add(Path(obsidian_dir).name)
        stem = record.get("obsidian_stem") or record.get("summary_stem")
        if stem:
            expected.add(stem)

    stale = [
        d for d in papers_root.iterdir()
        if d.is_dir() and not d.name.startswith((".", "_")) and d.name not in expected
    ]
    if not stale:
        return

    trash_dir = papers_root / "_trash" / time.strftime("%Y%m%d-%H%M%S")
    trash_dir.mkdir(parents=True, exist_ok=True)
    for folder in stale:
        try:
            shutil.move(str(folder), str(trash_dir / folder.name))
            log(f"[obsidian] moved stale folder to trash: {folder.name}")
        except Exception as exc:
            log(f"[obsidian] failed to move stale folder {folder.name}: {exc}")


def zotero_base_url(config: Config) -> str:
    return f"https://api.zotero.org/{config.zotero_library_type}/{config.zotero_user_id}"


def zotero_headers(config: Config, *, extra: dict[str, str] | None = None) -> dict[str, str]:
    api_key = os.environ.get(config.zotero_api_key_env, "").strip()
    if not api_key:
        raise SyncError(f"Missing environment variable {config.zotero_api_key_env}.")
    headers = {
        "Zotero-API-Key": api_key,
        "Accept": "application/json",
    }
    if extra:
        headers.update(extra)
    return headers


def zotero_get_item(config: Config, item_key: str) -> dict[str, Any]:
    _, _, payload = http_json(
        f"{zotero_base_url(config)}/items/{item_key}",
        headers=zotero_headers(config),
    )
    if not isinstance(payload, dict):
        raise SyncError(f"Unexpected Zotero item payload for {item_key}")
    return payload


def zotero_get_children(config: Config, item_key: str) -> list[dict[str, Any]]:
    _, _, payload = http_json(
        f"{zotero_base_url(config)}/items/{item_key}/children?format=json",
        headers=zotero_headers(config),
    )
    if not isinstance(payload, list):
        return []
    return [item for item in payload if isinstance(item, dict)]


def zotero_create_items(config: Config, items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    _, _, payload = http_json(
        f"{zotero_base_url(config)}/items",
        method="POST",
        headers=zotero_headers(config, extra={"Content-Type": "application/json"}),
        json_body=items,
    )
    if not isinstance(payload, dict):
        raise SyncError("Unexpected Zotero create-items response.")
    successful = payload.get("successful", {})
    if not isinstance(successful, dict) or not successful:
        raise SyncError(f"Zotero item creation failed: {json.dumps(payload, ensure_ascii=False)}")
    created: list[dict[str, Any]] = []
    for key in sorted(successful.keys()):
        value = successful[key]
        if isinstance(value, dict):
            created.append(value)
    return created


def zotero_put_item(config: Config, item_key: str, data: dict[str, Any]) -> None:
    http_json(
        f"{zotero_base_url(config)}/items/{item_key}",
        method="PUT",
        headers=zotero_headers(config, extra={"Content-Type": "application/json"}),
        json_body=data,
    )


def zotero_delete_item(config: Config, item_key: str, version: int | str | None = None) -> None:
    headers = zotero_headers(config)
    resolved_version = version
    if resolved_version is None:
        try:
            current = zotero_get_item(config, item_key)
            resolved_version = current.get("version") or current.get("data", {}).get("version")
        except Exception:
            resolved_version = None
    if resolved_version not in (None, ""):
        headers["If-Unmodified-Since-Version"] = str(resolved_version)
    http_json(
        f"{zotero_base_url(config)}/items/{item_key}",
        method="DELETE",
        headers=headers,
    )


def build_summary_note_html(record: dict[str, Any]) -> str:
    title = html.escape(record.get("title") or record.get("summary_stem") or "")
    summary_md = Path(record["summary_md"])
    package_dir = Path(record["package_dir"])
    obsidian_md = Path(record["obsidian_md"]) if record.get("obsidian_md") else None
    parts = [
        f"<p><strong>自动摘要同步</strong> <span data-codex-marker=\"{SUMMARY_NOTE_MARKER}\"></span></p>",
        "<ul>",
        f"<li>标题：{title}</li>",
        f"<li>本地总结：<a href=\"{html.escape(path_to_uri(summary_md))}\">{html.escape(summary_md.name)}</a></li>",
        f"<li>打包目录：<a href=\"{html.escape(path_to_uri(package_dir))}\">{html.escape(str(package_dir))}</a></li>",
    ]
    if obsidian_md and obsidian_md.exists():
        parts.append(
            f"<li>Obsidian 副本：<a href=\"{html.escape(path_to_uri(obsidian_md))}\">{html.escape(obsidian_md.name)}</a></li>"
        )
    deep_docx = Path(record["deep_package_docx"]) if record.get("deep_package_docx") else None
    deep_pdf = Path(record["deep_package_pdf"]) if record.get("deep_package_pdf") else None
    deep_dir = Path(record["deep_package_dir"]) if record.get("deep_package_dir") else None
    if (deep_docx and deep_docx.exists()) or (deep_pdf and deep_pdf.exists()) or (deep_dir and deep_dir.exists()):
        parts.append(f"<li><strong>精读 DOCX/PDF 包</strong> <span data-codex-marker=\"{DEEP_PACKAGE_NOTE_MARKER}\"></span><ul>")
        if deep_docx and deep_docx.exists():
            parts.append(
                f"<li>DOCX：<a href=\"{html.escape(path_to_uri(deep_docx))}\">{html.escape(deep_docx.name)}</a></li>"
            )
        if deep_pdf and deep_pdf.exists():
            parts.append(
                f"<li>PDF：<a href=\"{html.escape(path_to_uri(deep_pdf))}\">{html.escape(deep_pdf.name)}</a></li>"
            )
        if deep_dir and deep_dir.exists():
            parts.append(
                f"<li>目录：<a href=\"{html.escape(path_to_uri(deep_dir))}\">{html.escape(str(deep_dir))}</a></li>"
            )
        parts.append("</ul></li>")
    parts.append(f"<li>更新时间：{html.escape(record.get('generated_at', ''))}</li>")
    parts.append("</ul>")
    return "".join(parts)


def find_existing_summary_note(children: list[dict[str, Any]]) -> dict[str, Any] | None:
    for child in children:
        data = child.get("data", {})
        if data.get("itemType") == "note" and SUMMARY_NOTE_MARKER in first_nonempty(data.get("note")):
            return child
    return None


def find_existing_summary_attachment(children: list[dict[str, Any]], filename: str) -> dict[str, Any] | None:
    for child in children:
        data = child.get("data", {})
        if data.get("itemType") != "attachment":
            continue
        title = first_nonempty(data.get("title"))
        existing_filename = first_nonempty(data.get("filename"))
        path_value = first_nonempty(data.get("path"))
        path_name = Path(path_value).name if path_value else ""
        if (
            title.startswith(SUMMARY_ATTACHMENT_TITLE_PREFIX)
            or existing_filename == filename
            or path_name == filename
        ):
            return child
    return None


def ensure_zotero_summary_note(config: Config, parent_item_key: str, record: dict[str, Any]) -> None:
    children = zotero_get_children(config, parent_item_key)
    existing = find_existing_summary_note(children)
    note_html = build_summary_note_html(record)
    if existing is None:
        zotero_create_items(
            config,
            [{"itemType": "note", "parentItem": parent_item_key, "note": note_html}],
        )
        return

    current = zotero_get_item(config, existing["key"])
    data = current["data"]
    data["note"] = note_html
    zotero_put_item(config, existing["key"], data)


def authorize_zotero_upload(
    config: Config,
    attachment_key: str,
    *,
    filename: str,
    filesize: int,
    mtime_ms: int,
    md5_hash: str,
    previous_md5: str | None,
) -> dict[str, Any]:
    headers = zotero_headers(config)
    if previous_md5:
        headers["If-Match"] = previous_md5
    else:
        headers["If-None-Match"] = "*"
    _, _, payload = http_json(
        f"{zotero_base_url(config)}/items/{attachment_key}/file",
        method="POST",
        headers=headers,
        form_body={
            "md5": md5_hash,
            "filename": filename,
            "filesize": filesize,
            "mtime": mtime_ms,
            "params": 1,
        },
    )
    if not isinstance(payload, dict):
        raise SyncError("Unexpected Zotero upload authorization payload.")
    return payload


def upload_to_zotero_storage(auth_payload: dict[str, Any], file_path: Path) -> None:
    if auth_payload.get("exists") == 1:
        return
    params = auth_payload.get("params")
    if not isinstance(params, dict):
        raise SyncError(f"Unexpected Zotero upload auth payload: {json.dumps(auth_payload, ensure_ascii=False)}")
    body_bytes = file_path.read_bytes()
    content_type, body = multipart_form_data(
        {str(key): str(value) for key, value in params.items()},
        "file",
        file_path.name,
        body_bytes,
    )
    request = urllib.request.Request(
        auth_payload["url"],
        data=body,
        headers={"Content-Type": content_type},
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=600) as response:
            if response.status not in {200, 201, 204}:
                raise SyncError(f"Unexpected Zotero storage upload status: {response.status}")
    except urllib.error.HTTPError as exc:
        body_text = exc.read().decode("utf-8", errors="replace")
        raise SyncError(f"Zotero storage upload failed: HTTP {exc.code}: {body_text[:1000]}") from exc


def register_zotero_upload(config: Config, attachment_key: str, upload_key: str, previous_md5: str | None) -> None:
    headers = zotero_headers(config)
    if previous_md5:
        headers["If-Match"] = previous_md5
    else:
        headers["If-None-Match"] = "*"
    http_json(
        f"{zotero_base_url(config)}/items/{attachment_key}/file",
        method="POST",
        headers=headers,
        form_body={"upload": upload_key},
    )


def ensure_zotero_summary_attachment_imported_file(config: Config, parent_item_key: str, record: dict[str, Any]) -> None:
    summary_path = Path(record["summary_md"])
    children = zotero_get_children(config, parent_item_key)
    existing = find_existing_summary_attachment(children, summary_path.name)
    previous_md5: str | None = None
    previous_mtime: str | None = None
    attachment_key = ""
    delete_on_failure = False

    if existing is None:
        created = zotero_create_items(
            config,
            [
                {
                    "itemType": "attachment",
                    "parentItem": parent_item_key,
                    "linkMode": "imported_file",
                    "title": f"{SUMMARY_ATTACHMENT_TITLE_PREFIX}: {summary_path.name}",
                    "note": "",
                    "tags": [],
                    "relations": {},
                    "contentType": "text/markdown",
                    "charset": "utf-8",
                    "filename": summary_path.name,
                }
            ],
        )
        attachment_key = created[0]["key"]
        delete_on_failure = True
    else:
        attachment_key = existing["key"]
        current = zotero_get_item(config, attachment_key)
        data = current["data"]
        previous_md5 = first_nonempty(data.get("md5"))
        previous_mtime = first_nonempty(data.get("mtime"))
        data["title"] = f"{SUMMARY_ATTACHMENT_TITLE_PREFIX}: {summary_path.name}"
        data["contentType"] = "text/markdown"
        data["charset"] = "utf-8"
        data["filename"] = summary_path.name
        zotero_put_item(config, attachment_key, data)
        if not previous_md5 and not previous_mtime:
            delete_on_failure = True

    try:
        mtime_ms = int(summary_path.stat().st_mtime_ns // 1_000_000)
        auth_payload = authorize_zotero_upload(
            config,
            attachment_key,
            filename=summary_path.name,
            filesize=summary_path.stat().st_size,
            mtime_ms=mtime_ms,
            md5_hash=md5_file(summary_path),
            previous_md5=previous_md5,
        )
        if auth_payload.get("exists") == 1:
            return
        upload_to_zotero_storage(auth_payload, summary_path)
        register_zotero_upload(config, attachment_key, str(auth_payload["uploadKey"]), previous_md5)
    except Exception:
        if delete_on_failure and attachment_key:
            try:
                zotero_delete_item(config, attachment_key)
                log(f"[zotero] Removed incomplete summary attachment {attachment_key} after upload failure.")
            except Exception as cleanup_exc:
                log(f"[zotero] Failed to remove incomplete summary attachment {attachment_key}: {cleanup_exc}")
        raise


def ensure_zotero_summary_attachment_linked_file(config: Config, parent_item_key: str, record: dict[str, Any]) -> None:
    # Use a local cached copy of the packaged summary to avoid UNC + Unicode
    # path issues inside Zotero linked-file attachments while preserving
    # relative image references from the packaged folder.
    summary_path = prepare_zotero_linked_summary_path(record)
    children = zotero_get_children(config, parent_item_key)
    existing = find_existing_summary_attachment(children, summary_path.name)

    if existing is not None:
        current = zotero_get_item(config, existing["key"])
        data = current["data"]
        existing_mode = first_nonempty(data.get("linkMode")).lower()
        if existing_mode and existing_mode != "linked_file":
            zotero_delete_item(config, existing["key"], current.get("version"))
            existing = None
        else:
            data["title"] = f"{SUMMARY_ATTACHMENT_TITLE_PREFIX}: {summary_path.name}"
            data["contentType"] = "text/markdown"
            data["charset"] = "utf-8"
            data["path"] = str(summary_path)
            zotero_put_item(config, existing["key"], data)
            return

    zotero_create_items(
        config,
        [
            {
                "itemType": "attachment",
                "parentItem": parent_item_key,
                "linkMode": "linked_file",
                "title": f"{SUMMARY_ATTACHMENT_TITLE_PREFIX}: {summary_path.name}",
                "note": "",
                "tags": [],
                "relations": {},
                "contentType": "text/markdown",
                "charset": "utf-8",
                "path": str(summary_path),
            }
        ],
    )


def ensure_zotero_summary_attachment(config: Config, parent_item_key: str, record: dict[str, Any]) -> None:
    mode = config.zotero_summary_attachment_mode
    if mode == "none":
        return
    if mode == "linked_file":
        ensure_zotero_summary_attachment_linked_file(config, parent_item_key, record)
        return
    if mode == "imported_file":
        ensure_zotero_summary_attachment_imported_file(config, parent_item_key, record)
        return
    raise SyncError(f"Unsupported zotero_api.summary_attachment_mode: {mode}")


def ensure_zotero_linked_file_attachment(
    config: Config,
    parent_item_key: str,
    file_path: Path,
    *,
    title_prefix: str,
    content_type: str,
) -> None:
    if not file_path.exists():
        return
    children = zotero_get_children(config, parent_item_key)
    existing = find_existing_summary_attachment(children, file_path.name)
    title = f"{title_prefix}: {file_path.name}"
    if existing is not None:
        current = zotero_get_item(config, existing["key"])
        data = current["data"]
        existing_mode = first_nonempty(data.get("linkMode")).lower()
        if existing_mode and existing_mode != "linked_file":
            zotero_delete_item(config, existing["key"], current.get("version"))
            existing = None
        else:
            data["title"] = title
            data["contentType"] = content_type
            data["path"] = str(file_path.resolve())
            zotero_put_item(config, existing["key"], data)
            return
    zotero_create_items(
        config,
        [
            {
                "itemType": "attachment",
                "parentItem": parent_item_key,
                "linkMode": "linked_file",
                "title": title,
                "note": "",
                "tags": [],
                "relations": {},
                "contentType": content_type,
                "path": str(file_path.resolve()),
            }
        ],
    )


def ensure_zotero_deep_package_links(config: Config, parent_item_key: str, record: dict[str, Any]) -> None:
    docx = Path(record["deep_package_docx"]) if record.get("deep_package_docx") else None
    pdf = Path(record["deep_package_pdf"]) if record.get("deep_package_pdf") else None
    if docx and docx.exists():
        ensure_zotero_linked_file_attachment(
            config,
            parent_item_key,
            docx,
            title_prefix=DEEP_PACKAGE_ATTACHMENT_TITLE_PREFIX,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    if pdf and pdf.exists():
        ensure_zotero_linked_file_attachment(
            config,
            parent_item_key,
            pdf,
            title_prefix=DEEP_PACKAGE_ATTACHMENT_TITLE_PREFIX,
            content_type="application/pdf",
        )


def sync_zotero_item_tags(config: Config, item_key: str, ai_tags: list[str]) -> None:
    """Merge *ai_tags* into the Zotero item's existing tags.

    Each AI tag is added with ``type: 1`` (automatic/imported) so it can be
    distinguished from manually-added tags in the Zotero UI.
    Existing tags are preserved; duplicates are skipped.
    """
    if not ai_tags:
        return
    current = zotero_get_item(config, item_key)
    data = current["data"]
    existing_tag_names = {t["tag"] for t in data.get("tags", []) if isinstance(t, dict)}
    new_entries = [{"tag": t, "type": 1} for t in ai_tags if t not in existing_tag_names]
    if not new_entries:
        return
    data["tags"] = data.get("tags", []) + new_entries
    zotero_put_item(config, item_key, data)
    log(f"[zotero] Added {len(new_entries)} AI tags to {item_key}")


def sync_zotero_record(config: Config, record: dict[str, Any]) -> None:
    if not config.zotero_api_enabled:
        return
    if not config.zotero_user_id:
        log("[zotero] Skip write-back: zotero_api.user_id is empty.")
        return
    item_key = first_nonempty(record.get("zotero_item_key"))
    if not item_key:
        return
    log(f"[zotero] {record.get('summary_stem')} -> item {item_key}")
    if config.zotero_update_summary_note:
        ensure_zotero_summary_note(config, item_key, record)
    if config.zotero_attach_summary_markdown:
        try:
            ensure_zotero_summary_attachment(config, item_key, record)
        except Exception as exc:
            log(f"[zotero] Summary attachment sync skipped for {record.get('summary_stem')}: {exc}")
    if record.get("deep_package_docx") or record.get("deep_package_pdf"):
        try:
            ensure_zotero_deep_package_links(config, item_key, record)
        except Exception as exc:
            log(f"[zotero] Deep package link sync skipped for {record.get('summary_stem')}: {exc}")
    # Sync AI-generated tags from summary MD → Zotero item
    ai_tags = record.get("ai_tags", [])
    if ai_tags:
        try:
            sync_zotero_item_tags(config, item_key, ai_tags)
        except Exception as exc:
            log(f"[zotero] AI tag sync skipped for {record.get('summary_stem')}: {exc}")


def markdown_image_refs(text: str) -> list[str]:
    return [match.group(1).strip() for match in re.finditer(r"!\[[^\]]*\]\(([^)]+)\)", text)]


def markdown_image_entries(text: str, context_radius: int = 2) -> list[dict[str, Any]]:
    lines = text.splitlines()
    entries: list[dict[str, Any]] = []
    pattern = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
    for index, line in enumerate(lines):
        for match in pattern.finditer(line):
            start = max(0, index - context_radius)
            end = min(len(lines), index + context_radius + 1)
            entries.append(
                {
                    "line": index + 1,
                    "alt": match.group(1).strip(),
                    "ref": match.group(2).strip(),
                    "context": "\n".join(lines[start:end]).strip(),
                }
            )
    return entries


def markdown_plain_lines(text: str) -> list[tuple[str, str]]:
    lines: list[tuple[str, str]] = []
    in_code = False
    for raw in text.splitlines():
        stripped = raw.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            continue
        if in_code or not stripped or stripped.startswith("![](") or stripped.startswith("!["):
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            lines.append((f"h{len(heading.group(1))}", heading.group(2).strip()))
            continue
        cleaned = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", stripped)
        cleaned = re.sub(r"`([^`]+)`", r"\1", cleaned)
        cleaned = cleaned.strip()
        if cleaned:
            lines.append(("p", cleaned))
    return lines


def resolve_package_image(package_dir: Path, ref: str) -> Path | None:
    ref = urllib.parse.unquote(ref.strip())
    if ref.startswith(("http://", "https://", "file://")):
        return None
    candidate = (package_dir / ref).resolve()
    try:
        candidate.relative_to(package_dir.resolve())
    except ValueError:
        return None
    if candidate.is_file() and candidate.suffix.lower() in {".png", ".jpg", ".jpeg"}:
        return candidate
    return None


def image_dimensions_for_file(path: Path) -> tuple[int, int] | None:
    try:
        import fitz

        with fitz.open(str(path)) as image_doc:
            if image_doc.page_count < 1:
                return None
            rect = image_doc[0].rect
            return int(round(rect.width)), int(round(rect.height))
    except Exception:
        return None


def representative_package_page_dimensions(package_dir: Path) -> tuple[int, int] | None:
    candidates: list[Path] = []
    for pattern in (
        "pages/page-*.png",
        "extracted/*/pages/page-*.png",
        "extracted/**/pages/page-*.png",
    ):
        candidates.extend(sorted(package_dir.glob(pattern)))
    for candidate in candidates:
        size = image_dimensions_for_file(candidate)
        if size is not None:
            return size
    return None


def looks_like_full_page_image(
    image_size: tuple[int, int] | None,
    page_size: tuple[int, int] | None,
) -> bool:
    if image_size is None or page_size is None:
        return False
    image_width, image_height = image_size
    page_width, page_height = page_size
    if min(image_width, image_height, page_width, page_height) <= 0:
        return False
    width_ratio = image_width / page_width
    height_ratio = image_height / page_height
    image_ratio = image_width / max(image_height, 1)
    page_ratio = page_width / max(page_height, 1)
    area_ratio = (image_width * image_height) / max(page_width * page_height, 1)
    return width_ratio >= 0.82 and height_ratio >= 0.82 and area_ratio >= 0.70 and abs(image_ratio - page_ratio) <= 0.18


def build_image_reference_audit(record: dict[str, Any], summary_text: str) -> dict[str, Any]:
    summary_md = resolve_existing_path(record.get("summary_md"), kind="file")
    package_dir = resolve_existing_path(record.get("package_dir"), kind="dir")
    if package_dir is None and summary_md is not None:
        package_dir = summary_md.parent
    if package_dir is None:
        package_dir = Path(".").resolve()
    page_size = representative_package_page_dimensions(package_dir)
    entries: list[dict[str, Any]] = []
    issue_counts: dict[str, int] = {}
    for entry in markdown_image_entries(summary_text):
        ref = str(entry["ref"])
        resolved = resolve_package_image(package_dir, ref)
        issues: list[str] = []
        size: tuple[int, int] | None = None
        file_size = 0
        if resolved is None:
            issues.append("missing-or-external-image")
        else:
            size = image_dimensions_for_file(resolved)
            try:
                file_size = resolved.stat().st_size
            except OSError:
                file_size = 0
            if size is None:
                issues.append("unreadable-image")
            else:
                width, height = size
                if width < 140 or height < 90:
                    issues.append("too-small-likely-icon-or-fragment")
                if width / max(height, 1) > 8 or height / max(width, 1) > 8:
                    issues.append("extreme-aspect-ratio-possible-bad-crop")
                if looks_like_full_page_image(size, page_size):
                    issues.append("full-page-screenshot-possible-bad-figure-crop")
            if file_size and file_size < 3500:
                issues.append("very-small-file-possible-placeholder")
        for issue in issues:
            issue_counts[issue] = issue_counts.get(issue, 0) + 1
        entries.append(
            {
                "line": entry["line"],
                "alt": entry["alt"],
                "ref": ref,
                "resolved": str(resolved) if resolved is not None else "",
                "exists": bool(resolved and resolved.exists()),
                "size": list(size) if size is not None else None,
                "file_size": file_size,
                "issues": issues,
                "context": entry["context"],
            }
        )
    suspect_entries = [entry for entry in entries if entry["issues"]]
    return {
        "title": record.get("title") or record.get("summary_stem") or "",
        "summary_md": str(summary_md) if summary_md is not None else "",
        "package_dir": str(package_dir),
        "page_size_hint": list(page_size) if page_size is not None else None,
        "image_ref_count": len(entries),
        "suspect_count": len(suspect_entries),
        "issue_counts": issue_counts,
        "suspect_entries": suspect_entries[:80],
    }


def render_image_audit_report(audit: dict[str, Any], ai_report: str | None) -> str:
    lines = [
        "# Summary Image QA",
        "",
        f"- Title: {audit.get('title') or 'N/A'}",
        f"- Summary: `{audit.get('summary_md') or 'N/A'}`",
        f"- Package: `{audit.get('package_dir') or 'N/A'}`",
        f"- Image references checked: {audit.get('image_ref_count', 0)}",
        f"- Suspect references: {audit.get('suspect_count', 0)}",
        "",
    ]
    if not audit.get("suspect_count"):
        lines.extend(
            [
                "## 结论",
                "",
                "本地图片 QA 通过：Markdown 中的图片引用均可解析，未发现缺失图片、疑似整页截图、过小碎片图或极端比例裁剪风险。本次未调用 AI，以节省模型资源。",
                "",
            ]
        )
        return "\n".join(lines).rstrip() + "\n"

    lines.extend(["## 本地规则发现的问题", ""])
    issue_counts = audit.get("issue_counts", {})
    if isinstance(issue_counts, dict):
        for issue, count in sorted(issue_counts.items()):
            lines.append(f"- {issue}: {count}")
    lines.append("")
    lines.append("## 可疑图片引用")
    lines.append("")
    for entry in audit.get("suspect_entries", []):
        if not isinstance(entry, dict):
            continue
        lines.append(f"### Line {entry.get('line')}: `{entry.get('ref')}`")
        lines.append(f"- Issues: {', '.join(entry.get('issues') or [])}")
        lines.append(f"- Size: {entry.get('size') or 'N/A'}")
        lines.append(f"- Exists: {entry.get('exists')}")
        if entry.get("context"):
            lines.append("")
            lines.append("```markdown")
            lines.append(str(entry.get("context")))
            lines.append("```")
        lines.append("")
    if ai_report:
        lines.extend(["## AI 判断", "", ai_report.rstrip(), ""])
    return "\n".join(lines).rstrip() + "\n"


def collect_deep_package_images(summary_text: str, package_dir: Path, limit: int = 12) -> list[Path]:
    images: list[Path] = []
    seen: set[str] = set()
    for ref in markdown_image_refs(summary_text):
        image_path = resolve_package_image(package_dir, ref)
        if image_path is None:
            continue
        key = str(image_path)
        if key in seen:
            continue
        seen.add(key)
        images.append(image_path)
        if len(images) >= limit:
            break
    images_dir = package_dir / "images"
    if images_dir.is_dir() and len(images) < limit:
        for image_path in sorted(images_dir.glob("*.png")):
            key = str(image_path.resolve())
            if key in seen:
                continue
            seen.add(key)
            images.append(image_path.resolve())
            if len(images) >= limit:
                break
    return images


def deep_package_dir_for_record(record: dict[str, Any]) -> Path:
    package_dir = Path(first_nonempty(record.get("package_dir")))
    if package_dir:
        return package_dir / "deep_reading_package"
    summary_md = Path(first_nonempty(record.get("summary_md"), "summary.md"))
    return summary_md.parent / f"{summary_md.stem}_deep_reading_package"


def write_deep_docx(record: dict[str, Any], summary_text: str, images: list[Path], output_path: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except Exception as exc:
        raise SyncError(
            "python-docx is required for deep DOCX package generation. "
            "Run setup_windows.ps1 or install requirements.txt in the pdf_tools venv."
        ) from exc

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    title = record.get("title") or record.get("summary_stem") or "Deep Reading Package"
    document.add_heading(str(title), level=0)
    meta = document.add_paragraph()
    meta.add_run("Deep reading package generated by local Paper Reading Workflow.").italic = True
    info = document.add_table(rows=0, cols=2)
    info.style = "Table Grid"
    for label, value in (
        ("Year", record.get("year", "")),
        ("Source", record.get("source", "")),
        ("Authors", ", ".join(record.get("authors", []) if isinstance(record.get("authors"), list) else [])),
        ("DOI", record.get("doi", "")),
        ("arXiv", record.get("arxiv", "")),
        ("PDF", record.get("pdf_path", "")),
    ):
        if not value:
            continue
        row = info.add_row().cells
        row[0].text = label
        row[1].text = str(value)

    document.add_heading("Reading Notes", level=1)
    for kind, value in markdown_plain_lines(summary_text):
        if kind == "h1":
            document.add_heading(value, level=1)
        elif kind in {"h2", "h3"}:
            document.add_heading(value, level=2 if kind == "h2" else 3)
        elif value.startswith(("- ", "* ")):
            document.add_paragraph(value[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s+", value):
            document.add_paragraph(re.sub(r"^\d+\.\s+", "", value), style="List Number")
        else:
            document.add_paragraph(value)

    if images:
        document.add_heading("Key Figures And Formula Snapshots", level=1)
        for image_path in images:
            document.add_paragraph(image_path.name)
            try:
                document.add_picture(str(image_path), width=Inches(5.8))
            except Exception as exc:
                document.add_paragraph(f"[image skipped: {image_path.name}: {exc}]")

    ensure_dir(output_path.parent)
    document.save(str(output_path))


def write_deep_pdf(record: dict[str, Any], summary_text: str, images: list[Path], output_path: Path) -> None:
    try:
        import fitz
    except Exception as exc:
        raise SyncError("PyMuPDF is required for deep PDF package generation.") from exc

    doc = fitz.open()
    width, height = fitz.paper_size("a4")
    margin = 50
    font = "china-s"

    def new_page() -> tuple[fitz.Page, float]:
        page = doc.new_page(width=width, height=height)
        return page, margin

    page, y = new_page()
    title = str(record.get("title") or record.get("summary_stem") or "Deep Reading Package")
    y += page.insert_textbox(fitz.Rect(margin, y, width - margin, y + 64), title, fontsize=18, fontname=font)
    y += 16
    metadata_lines = [
        f"Year: {record.get('year', '')}",
        f"Source: {record.get('source', '')}",
        f"Authors: {', '.join(record.get('authors', []) if isinstance(record.get('authors'), list) else [])}",
        f"PDF: {record.get('pdf_path', '')}",
    ]
    for line in metadata_lines:
        if line.endswith(": "):
            continue
        y += page.insert_textbox(fitz.Rect(margin, y, width - margin, y + 38), line, fontsize=9, fontname=font)
        y += 4

    for kind, value in markdown_plain_lines(summary_text):
        font_size = 13 if kind in {"h1", "h2"} else 10
        box_height = 54 if kind in {"h1", "h2"} else 42
        if y + box_height > height - margin:
            page, y = new_page()
        text = value[:1400]
        y += page.insert_textbox(fitz.Rect(margin, y, width - margin, y + box_height), text, fontsize=font_size, fontname=font)
        y += 6

    for image_path in images:
        if y + 230 > height - margin:
            page, y = new_page()
        y += page.insert_textbox(fitz.Rect(margin, y, width - margin, y + 22), image_path.name, fontsize=9, fontname=font)
        y += 4
        try:
            pix = fitz.Pixmap(str(image_path))
            img_w = width - 2 * margin
            img_h = img_w * pix.height / max(pix.width, 1)
            max_h = min(360, height - margin - y)
            if img_h > max_h:
                img_h = max_h
                img_w = img_h * pix.width / max(pix.height, 1)
            page.insert_image(fitz.Rect(margin, y, margin + img_w, y + img_h), filename=str(image_path))
            y += img_h + 14
        except Exception as exc:
            y += page.insert_textbox(fitz.Rect(margin, y, width - margin, y + 32), f"[image skipped: {exc}]", fontsize=9, fontname=font)
            y += 8

    ensure_dir(output_path.parent)
    doc.save(str(output_path))
    doc.close()


def generate_deep_reading_package(config: Config, record: dict[str, Any]) -> dict[str, Any]:
    summary_md = resolve_existing_path(record.get("summary_md"), kind="file")
    package_dir = resolve_existing_path(record.get("package_dir"), kind="dir")
    if summary_md is None:
        raise SyncError("Existing summary markdown is required before generating a deep package.")
    if package_dir is None:
        package_dir = summary_md.parent
    summary_text = read_text(summary_md)
    output_dir = deep_package_dir_for_record(record)
    ensure_dir(output_dir)
    stem = sanitize_windows_filename(record.get("summary_stem") or summary_md.stem, summary_md.stem, max_length=96)
    images = collect_deep_package_images(summary_text, package_dir)
    docx_path = output_dir / f"{stem}-deep-reading.docx"
    pdf_path = output_dir / f"{stem}-deep-reading.pdf"
    manifest_path = output_dir / "deep_reading_manifest.json"
    log(f"[deep-package] {summary_md.name}")
    write_deep_docx(record, summary_text, images, docx_path)
    write_deep_pdf(record, summary_text, images, pdf_path)
    manifest = {
        "summary_md": str(summary_md.resolve()),
        "package_dir": str(package_dir.resolve()),
        "docx": str(docx_path.resolve()),
        "pdf": str(pdf_path.resolve()),
        "images": [str(path) for path in images],
        "generated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
    }
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    updated = dict(record)
    updated["deep_package_dir"] = str(output_dir.resolve())
    updated["deep_package_docx"] = str(docx_path.resolve())
    updated["deep_package_pdf"] = str(pdf_path.resolve())
    updated["deep_package_manifest"] = str(manifest_path.resolve())
    updated["deep_package_generated_at"] = manifest["generated_at"]
    return updated


def enrich_record_with_match(
    config: Config,
    pdf_path: Path,
    inferred: dict[str, Any],
    matched: dict[str, Any] | None,
    previous: dict[str, Any] | None,
) -> dict[str, Any]:
    title = matched["title"] if matched and matched.get("title") else inferred.get("title") or (previous or {}).get("title") or pdf_path.stem
    year = matched["year"] if matched and matched.get("year") else (previous or {}).get("year", "")
    source = matched["source"] if matched and matched.get("source") else (previous or {}).get("source", "")
    tags = matched["tags"] if matched else (previous or {}).get("tags", [])
    creators = matched["creators"] if matched and matched.get("creators") else (previous or {}).get("creators", [])
    authors = display_authors_from_creators(creators)
    obsidian_tags = [f"{config.obsidian_tag_prefix}/{sanitize_obsidian_tag(tag)}" for tag in tags]
    return {
        "title": title,
        "year": year,
        "source": source,
        "tags": tags,
        "creators": creators,
        "authors": authors,
        "obsidian_tags": sorted(dict.fromkeys(obsidian_tags)),
        "doi": matched["doi"] if matched else inferred.get("doi", ""),
        "arxiv": matched["arxiv"] if matched else inferred.get("arxiv", ""),
        "citation_key": matched["citation_key"] if matched else (previous or {}).get("citation_key", ""),
        "zotero_item_key": matched["item_key"] if matched else (previous or {}).get("zotero_item_key", ""),
    }


def compare_spec_path(config: Config) -> Path:
    return config.workspace_root / "02-paper_summary_specs" / "compare.md"


def comparison_root(config: Config) -> Path:
    return config.summary_root / "06-paper_comparisons"


def derive_comparison_stem(records: list[dict[str, Any]]) -> str:
    year_values = [str(record.get("year") or "")[:4] for record in records if str(record.get("year") or "")[:4]]
    year = max(year_values) if year_values else str(time.localtime().tm_year)
    authors: list[str] = []
    title_words: list[str] = []
    for record in records:
        record_authors = record.get("authors", [])
        if isinstance(record_authors, list) and record_authors:
            authors.append(filename_fragment(str(record_authors[0]), 2, 18) or "Paper")
        else:
            authors.append(filename_fragment(str(record.get("title") or "Paper"), 2, 18) or "Paper")
        title_words.extend(re.findall(r"[A-Za-z0-9]+", str(record.get("title") or "")))
    topic = filename_fragment(" ".join(title_words[:12]), 4, 32) or "Multi-paper-comparison"
    author_part = "-vs-".join(author[:18] for author in authors[:4])
    stem = f"Compare-{year}-{topic}-{author_part}".strip("-")
    return stem[:120].rstrip("-") or f"Compare-{year}-Multi-paper-comparison"


def build_comparison_prompt(
    config: Config,
    pdf_entries: list[dict[str, Any]],
    spec_text: str,
) -> list[dict[str, Any]]:
    paper_payloads: list[dict[str, Any]] = []
    remaining_chars = max(config.openai_max_input_chars, 100000)
    per_paper_limit = max(30000, remaining_chars // max(len(pdf_entries), 1))
    for index, entry in enumerate(pdf_entries, start=1):
        bundle = entry["bundle"]
        extracted_md = str(bundle.get("document_md") or "")
        truncated = False
        if len(extracted_md) > per_paper_limit:
            keep_head = int(per_paper_limit * 0.65)
            keep_tail = per_paper_limit - keep_head
            extracted_md = (
                extracted_md[:keep_head]
                + "\n\n[... middle omitted because the comparison input is long ...]\n\n"
                + extracted_md[-keep_tail:]
            )
            truncated = True
        paper_payloads.append(
            {
                "index": index,
                "pdf_path": str(entry["pdf_path"].resolve()),
                "inferred_metadata": entry["inferred"],
                "matched_zotero_metadata": entry.get("matched"),
                "existing_summary_md": entry.get("existing_summary_md", ""),
                "extraction_dir": str(bundle["outdir"]),
                "input_truncated": truncated,
                "extracted_markdown": extracted_md,
            }
        )

    system_prompt = (
        "你要为本地论文管理工作流生成一份中文 Markdown 多篇文献对比总结。"
        "输出只能是最终 Markdown，不要加解释，不要加代码围栏。"
        "必须严格区分原文证据、跨文献比较和基于背景知识的补充判断。"
    )
    user_prompt = (
        "请按照下面的多篇文献对比规范，生成一份完整中文 Markdown 对比总结。\n\n"
        f"多篇对比规范全文：\n{spec_text}\n\n"
        f"本次用户补充要求和关键问题：\n{config.summary_user_request or '（无）'}\n\n"
        "参与对比的论文材料如下。每篇论文的 extracted_markdown 来自 PDF 自动提取，"
        "其中的本地图片路径可以在相关章节引用；如无法可靠定位图片，打包阶段会补充自动提取图片附录。\n\n"
        f"{json.dumps(paper_payloads, ensure_ascii=False, indent=2)}\n"
    )
    return [
        {"role": "system", "content": [{"type": "input_text", "text": system_prompt}]},
        {"role": "user", "content": [{"type": "input_text", "text": user_prompt}]},
    ]


def upsert_related_comparison_link(md_path: Path, comparison_title: str, comparison_md: Path) -> None:
    if not md_path.exists():
        return
    text = read_text(md_path)
    marker = "<!-- codex-related-comparisons -->"
    link = f"- [{comparison_title}]({path_to_uri(comparison_md)})"
    if str(comparison_md) in text or path_to_uri(comparison_md) in text:
        return
    section = f"\n\n## Related Comparison Summaries\n\n{marker}\n{link}\n"
    if marker not in text:
        md_path.write_text(text.rstrip() + section, encoding="utf-8")
        return
    updated = text.replace(marker, marker + "\n" + link, 1)
    md_path.write_text(updated, encoding="utf-8")


def sync_zotero_comparison_note(config: Config, parent_item_key: str, comparison_record: dict[str, Any]) -> None:
    if not config.zotero_api_enabled or not parent_item_key:
        return
    marker = f"codex-zotero-comparison-note-{comparison_record['summary_stem']}"
    children = zotero_get_children(config, parent_item_key)
    existing = None
    for child in children:
        data = child.get("data", {})
        if data.get("itemType") == "note" and marker in first_nonempty(data.get("note")):
            existing = child
            break
    summary_md = Path(comparison_record["summary_md"])
    package_dir = Path(comparison_record["package_dir"])
    note_html = (
        f"<p><strong>AI Comparison Summary (Codex)</strong> "
        f"<span data-codex-marker=\"{marker}\"></span></p>"
        "<ul>"
        f"<li>Comparison: {html.escape(comparison_record.get('title') or comparison_record['summary_stem'])}</li>"
        f"<li>Markdown: <a href=\"{html.escape(path_to_uri(summary_md))}\">{html.escape(summary_md.name)}</a></li>"
        f"<li>Package: <a href=\"{html.escape(path_to_uri(package_dir))}\">{html.escape(str(package_dir))}</a></li>"
        "</ul>"
    )
    if existing is None:
        zotero_create_items(config, [{"itemType": "note", "parentItem": parent_item_key, "note": note_html}])
        return
    current = zotero_get_item(config, existing["key"])
    data = current["data"]
    data["note"] = note_html
    zotero_put_item(config, existing["key"], data)


def run_compare(config: Config, explicit_pdfs: list[Path] | None = None) -> int:
    if not explicit_pdfs or len(explicit_pdfs) < 2:
        raise SyncError("The compare command requires at least two --pdf paths.")
    ensure_dir(config.cache_root)
    ensure_dir(comparison_root(config))

    spec_path = compare_spec_path(config)
    if not spec_path.exists():
        raise FileNotFoundError(f"Compare spec file not found: {spec_path}")
    spec_text = read_text(spec_path)

    export_records = load_export_records(config.export_path)
    if not export_records:
        export_records = load_export_records_from_local_zotero(config)
    export_index = build_export_index(export_records)
    state = load_json(config.state_path, {"spec_hash": "", "items": {}})
    items_state = state.get("items", {}) if isinstance(state.get("items"), dict) else {}

    pdf_entries: list[dict[str, Any]] = []
    related_records: list[dict[str, Any]] = []
    for pdf_path in discover_pdfs(config, explicit_pdfs):
        bundle = load_extraction_bundle(extract_pdf(config, pdf_path, force=False))
        inferred = infer_pdf_metadata(pdf_path, bundle)
        matched, match_score, match_reason = match_export_record(pdf_path, inferred, export_index)
        previous = items_state.get(str(pdf_path.resolve()))
        metadata = enrich_record_with_match(config, pdf_path, inferred, matched, previous if isinstance(previous, dict) else None)
        existing_summary = ""
        if isinstance(previous, dict):
            existing_path = resolve_existing_path(previous.get("summary_md"), kind="file")
            existing_summary = str(existing_path) if existing_path else ""
        entry_record = {
            "pdf_path": str(pdf_path.resolve()),
            "summary_md": existing_summary,
            "summary_stem": Path(existing_summary).stem if existing_summary else "",
            "match_score": round(match_score, 4),
            "match_reason": match_reason,
            **metadata,
        }
        pdf_entries.append(
            {
                "pdf_path": pdf_path,
                "bundle": bundle,
                "inferred": inferred,
                "matched": matched,
                "existing_summary_md": existing_summary,
                "record": entry_record,
            }
        )
        related_records.append(entry_record)

    stem = derive_comparison_stem(related_records)
    compare_dir = comparison_root(config) / stem
    summary_md = comparison_root(config) / f"{stem}.md"
    prompt = build_comparison_prompt(config, pdf_entries, spec_text)
    log(f"[summarize] compare {len(pdf_entries)} PDFs -> {summary_md.name}")
    summary_md.write_text(call_model(config, prompt), encoding="utf-8")
    package_dir = package_summary(
        config,
        summary_md,
        [entry["pdf_path"] for entry in pdf_entries],
        clean=True,
        package_dir=compare_dir,
        spec_path=spec_path,
    )
    related_path = package_dir / "related_papers.json"
    related_path.write_text(json.dumps(related_records, ensure_ascii=False, indent=2), encoding="utf-8")

    comparison_record = {
        "title": summary_md.stem,
        "summary_md": str((package_dir / summary_md.name).resolve()),
        "package_dir": str(package_dir.resolve()),
        "summary_stem": stem,
        "generated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "pdf_path": str(pdf_entries[0]["pdf_path"].resolve()),
        "pdf_paths": [str(entry["pdf_path"].resolve()) for entry in pdf_entries],
        "authors": [],
        "tags": ["comparison"],
        "obsidian_tags": [f"{config.obsidian_tag_prefix}/comparison"],
        "year": time.strftime("%Y"),
        "source": "multi-paper comparison",
        "zotero_item_key": "",
        "citation_key": "",
        "doi": "",
        "arxiv": "",
        "related_papers": related_records,
    }
    comparison_record["obsidian_stem"] = stem
    comparison_record = sync_obsidian_package(config, comparison_record, copy_package=True)

    for entry in pdf_entries:
        record = entry["record"]
        comparison_md = Path(comparison_record["summary_md"])
        if record.get("summary_md"):
            upsert_related_comparison_link(Path(record["summary_md"]), stem, comparison_md)
        previous = items_state.get(str(entry["pdf_path"].resolve()))
        if isinstance(previous, dict) and previous.get("obsidian_md"):
            upsert_related_comparison_link(Path(previous["obsidian_md"]), stem, comparison_md)
        item_key = first_nonempty(record.get("zotero_item_key"))
        if item_key:
            try:
                sync_zotero_comparison_note(config, item_key, comparison_record)
            except Exception as exc:
                log(f"[zotero] comparison note skipped for {item_key}: {exc}")

    items_state[f"comparison::{stem}"] = comparison_record
    all_records = [record for record in items_state.values() if isinstance(record, dict)]
    rebuild_obsidian_indexes(config, all_records)
    save_json(
        config.state_path,
        {
            **state,
            "items": items_state,
            "last_compare_run": time.strftime("%Y-%m-%d %H:%M:%S"),
        },
    )
    log("[done] Compare pass completed.")
    return 0


def process_pdf(
    config: Config,
    pdf_path: Path,
    *,
    state_record: dict[str, Any] | None,
    export_index: dict[str, dict[str, list[dict[str, Any]]]],
    spec_text: str,
    spec_changed: bool,
) -> dict[str, Any]:
    pdf_sig = file_signature(pdf_path)
    pdf_changed = state_record is None or state_record.get("pdf_sig") != pdf_sig
    summary_missing = not state_record or resolve_existing_path(state_record.get("summary_md"), kind="file") is None
    package_missing = not state_record or resolve_existing_path(state_record.get("package_dir"), kind="dir") is None
    force_regenerate = spec_changed or pdf_changed or summary_missing

    bundle = load_extraction_bundle(extract_pdf(config, pdf_path, force=force_regenerate))
    inferred = infer_pdf_metadata(pdf_path, bundle)
    matched, match_score, match_reason = match_export_record(pdf_path, inferred, export_index)
    metadata = enrich_record_with_match(config, pdf_path, inferred, matched, state_record)

    summary_md, package_dir = ensure_summary_and_package(
        config,
        pdf_path,
        inferred,
        matched,
        bundle,
        spec_text,
        force_regenerate=force_regenerate or package_missing,
        existing_record=state_record if not (spec_changed or pdf_changed) else None,
    )

    record = {
        "pdf_path": str(pdf_path.resolve()),
        "pdf_sig": pdf_sig,
        "pdf_sha256": sha256_file(pdf_path),
        "summary_md": str(summary_md.resolve()),
        "package_dir": str(package_dir.resolve()),
        "summary_stem": summary_md.stem,
        "generated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "match_score": round(match_score, 4),
        "match_reason": match_reason,
        "inferred_title": inferred.get("title", ""),
        **metadata,
    }
    record["obsidian_stem"] = derive_obsidian_item_stem(record)

    obsidian_needs_copy = (
        force_regenerate
        or package_missing
        or not state_record
        or resolve_existing_path(state_record.get("obsidian_dir"), kind="dir") is None
    )
    record = sync_obsidian_package(config, record, copy_package=obsidian_needs_copy)
    sync_zotero_record(config, record)
    return record


def discover_pdfs(config: Config, explicit: list[Path] | None = None) -> list[Path]:
    if explicit:
        resolved: list[Path] = []
        for path in explicit:
            candidate = path.expanduser()
            if not candidate.is_absolute():
                candidate = (config.workspace_root / candidate).resolve()
            else:
                candidate = candidate.resolve()
            resolved.append(candidate)
        return resolved
    return sorted(config.paper_dir.rglob("*.pdf"))


def run_once(config: Config, explicit_pdfs: list[Path] | None = None, *, force_regenerate: bool = False) -> int:
    ensure_dir(config.cache_root)
    ensure_dir(config.export_path.parent)
    ensure_dir(config.obsidian_vault_dir)
    ensure_dir(config.obsidian_vault_dir / config.obsidian_papers_subdir)
    ensure_dir(config.obsidian_vault_dir / config.obsidian_tags_subdir)

    if not config.extract_script.exists() or not config.package_script.exists():
        raise FileNotFoundError("Could not find the existing PDF extraction / packaging scripts under 03-tools.")
    if not config.spec_path.exists():
        raise FileNotFoundError(f"Spec file not found: {config.spec_path}")

    state = load_json(config.state_path, {"spec_hash": "", "items": {}})
    items_state = state.get("items", {})
    if not isinstance(items_state, dict):
        items_state = {}

    spec_text = read_text(config.spec_path)
    spec_hash = sha256_text(spec_text)
    spec_changed = state.get("spec_hash") != spec_hash
    if spec_changed:
        log("[spec] default.md changed, summaries will be regenerated with the latest spec.")

    export_records = load_export_records(config.export_path)
    if export_records:
        log(f"[match] Loaded {len(export_records)} records from BBT export.")
    else:
        log("[match] BBT export missing or empty; falling back to local Zotero index.")
        export_records = load_export_records_from_local_zotero(config)
    export_index = build_export_index(export_records)
    pdfs = discover_pdfs(config, explicit_pdfs)
    log(f"[scan] PDFs found: {len(pdfs)}")

    incremental_run = bool(explicit_pdfs)
    new_state_items: dict[str, Any] = dict(items_state) if incremental_run else {}
    processed_records: list[dict[str, Any]] = []

    for pdf_path in pdfs:
        key = str(pdf_path.resolve())
        previous = items_state.get(key) if isinstance(items_state.get(key), dict) else None
        try:
            record = process_pdf(
                config,
                pdf_path,
                state_record=previous,
                export_index=export_index,
                spec_text=spec_text,
                spec_changed=spec_changed or force_regenerate,
            )
            new_state_items[key] = record
            processed_records.append(record)
        except Exception as exc:
            log(f"[error] {pdf_path.name}: {exc}")
            if previous:
                previous["last_error"] = str(exc)
                new_state_items[key] = previous
                processed_records.append(previous)

    all_records = [record for record in new_state_items.values() if isinstance(record, dict)]
    rebuild_obsidian_indexes(config, all_records)
    if not incremental_run:
        cleanup_stale_obsidian_folders(config, all_records)
    save_json(
        config.state_path,
        {
            "spec_hash": spec_hash,
            "items": new_state_items,
            "last_run": time.strftime("%Y-%m-%d %H:%M:%S"),
        },
    )
    log("[done] Sync pass completed.")
    return 0


def _state_and_items(config: Config) -> tuple[dict[str, Any], dict[str, Any]]:
    state = load_json(config.state_path, {"spec_hash": "", "items": {}})
    items_state = state.get("items", {}) if isinstance(state.get("items"), dict) else {}
    return state, items_state


def run_deep_package(config: Config, explicit_pdfs: list[Path] | None = None) -> int:
    if not explicit_pdfs:
        raise SyncError("The deep-package command requires at least one --pdf path.")
    state, items_state = _state_and_items(config)
    updated_records: list[dict[str, Any]] = []
    for pdf_path in discover_pdfs(config, explicit_pdfs):
        key = str(pdf_path.resolve())
        previous = items_state.get(key) if isinstance(items_state.get(key), dict) else None
        if previous is None:
            raise SyncError(f"No existing summary state found for {pdf_path.name}. Generate a summary first.")
        package_missing = resolve_existing_path(previous.get("package_dir"), kind="dir") is None
        summary_md = resolve_existing_path(previous.get("summary_md"), kind="file")
        if summary_md is None:
            raise SyncError(f"No existing summary markdown found for {pdf_path.name}.")
        if package_missing:
            package_dir = package_summary(config, summary_md, pdf_path, clean=True)
            previous["package_dir"] = str(package_dir.resolve())
        record = generate_deep_reading_package(config, previous)
        record = sync_obsidian_package(config, record, copy_package=True)
        items_state[key] = record
        updated_records.append(record)
        sync_zotero_record(config, record)

    all_records = [record for record in items_state.values() if isinstance(record, dict)]
    rebuild_obsidian_indexes(config, all_records)
    save_json(
        config.state_path,
        {
            **state,
            "items": items_state,
            "last_deep_package_run": time.strftime("%Y-%m-%d %H:%M:%S"),
        },
    )
    log(f"[done] Deep reading package generated for {len(updated_records)} item(s).")
    return 0


def build_summary_quality_prompt(config: Config, record: dict[str, Any], image_audit: dict[str, Any]) -> list[dict[str, Any]]:
    system_prompt = (
        "You are a strict but low-cost QA reviewer for a paper-reading workflow. "
        "Focus only on Markdown image references, figure/formula crop risks, and whether the cited image likely matches its surrounding text. "
        "Return only a concise Markdown report in Chinese. Do not review the whole paper summary."
    )
    user_prompt = (
        "请基于下面的本地图像 QA 结果，重点判断 Markdown 总结中的图片引用和裁剪是否可能有问题。\n"
        "本地规则已经检查了文件是否存在、图片尺寸、疑似整页截图、过小碎片图和极端比例。"
        "你只需要复核这些可疑项，给出是否需要人工打开核查、建议如何修复。"
        "不要做全文规范性审稿，不要讨论与图片无关的内容，以节省资源。\n\n"
        f"论文元数据：\n{json.dumps({k: record.get(k) for k in ['title', 'year', 'source', 'authors', 'doi', 'arxiv', 'pdf_path']}, ensure_ascii=False, indent=2)}\n\n"
        f"图像 QA 结果：\n{json.dumps(image_audit, ensure_ascii=False, indent=2)}\n\n"
        "输出格式必须简洁：\n"
        "## 结论\n- 通过/需修改\n\n"
        "## 图片问题\n- [严重程度] 行号、图片路径、问题、建议\n\n"
        "## 人工复核优先级\n- ...\n"
    )
    return [
        {"role": "system", "content": [{"type": "input_text", "text": system_prompt}]},
        {"role": "user", "content": [{"type": "input_text", "text": user_prompt}]},
    ]


def run_check_summary(config: Config, explicit_pdfs: list[Path] | None = None) -> int:
    if not explicit_pdfs:
        raise SyncError("The check-summary command requires at least one --pdf path.")
    if not config.openai_enabled and not codex_cli_available(config):
        raise SyncError("No enabled model backend is available for summary QA.")
    state, items_state = _state_and_items(config)
    checked = 0
    for pdf_path in discover_pdfs(config, explicit_pdfs):
        key = str(pdf_path.resolve())
        record = items_state.get(key) if isinstance(items_state.get(key), dict) else None
        if record is None:
            raise SyncError(f"No existing summary state found for {pdf_path.name}.")
        summary_md = resolve_existing_path(record.get("summary_md"), kind="file")
        if summary_md is None:
            raise SyncError(f"No existing summary markdown found for {pdf_path.name}.")
        summary_text = read_text(summary_md)
        log(f"[check] image refs in {summary_md.name}")
        image_audit = build_image_reference_audit(record, summary_text)
        ai_report = None
        if image_audit.get("suspect_count"):
            log(f"[check] AI image QA for {summary_md.name}: {image_audit.get('suspect_count')} suspect ref(s)")
            ai_report = call_model(config, build_summary_quality_prompt(config, record, image_audit))
        else:
            log(f"[check] local image QA passed for {summary_md.name}; skip AI call")
        report = render_image_audit_report(image_audit, ai_report)
        output_dir = deep_package_dir_for_record(record)
        ensure_dir(output_dir)
        report_path = output_dir / f"{summary_md.stem}-image-quality-check.md"
        report_path.write_text(report, encoding="utf-8")
        record["summary_quality_check"] = str(report_path.resolve())
        record["summary_quality_checked_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
        record["summary_quality_check_focus"] = "image-references-and-crops"
        items_state[key] = record
        checked += 1

    save_json(
        config.state_path,
        {
            **state,
            "items": items_state,
            "last_summary_check_run": time.strftime("%Y-%m-%d %H:%M:%S"),
        },
    )
    log(f"[done] Summary QA completed for {checked} item(s).")
    return 0


def call_bbt_jsonrpc(config: Config, method: str, params: list[Any]) -> Any:
    payload = {"jsonrpc": "2.0", "method": method, "params": params, "id": 1}
    _, _, response = http_json(
        config.bbt_jsonrpc_url,
        method="POST",
        headers={"Content-Type": "application/json", "Accept": "application/json"},
        json_body=payload,
        timeout=60,
    )
    if isinstance(response, dict) and response.get("error"):
        raise SyncError(f"BBT JSON-RPC error: {json.dumps(response['error'], ensure_ascii=False)}")
    if isinstance(response, dict) and "result" in response:
        return response["result"]
    raise SyncError(f"Unexpected BBT JSON-RPC response: {json.dumps(response, ensure_ascii=False)}")


def setup_bbt(config: Config) -> int:
    ensure_dir(config.export_path.parent)
    ready = call_bbt_jsonrpc(config, "api.ready", [])
    result = call_bbt_jsonrpc(
        config,
        "autoexport.add",
        [
            config.bbt_collection_path,
            config.bbt_translator,
            str(config.export_path),
            config.bbt_display_options,
            config.bbt_replace,
        ],
    )
    log(
        json.dumps(
            {
                "bbt_ready": ready,
                "autoexport": result,
                "export_path": str(config.export_path),
                "collection_path": config.bbt_collection_path,
                "translator": config.bbt_translator,
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0


def watch(config: Config, interval_override: int | None) -> int:
    interval = interval_override or config.poll_interval_seconds
    if interval <= 0:
        interval = 30
    log(f"[watch] Polling every {interval} second(s).")
    while True:
        try:
            run_once(config)
        except KeyboardInterrupt:
            raise
        except Exception as exc:
            log(f"[watch-error] {exc}")
        time.sleep(interval)


def main() -> int:
    args = parse_args()
    config = apply_runtime_model_overrides(load_config(args.config), args)
    if args.command == "setup-bbt":
        return setup_bbt(config)
    if args.command == "watch":
        return watch(config, args.interval)
    if args.command == "compare":
        return run_compare(config, args.pdf)
    if args.command == "deep-package":
        return run_deep_package(config, args.pdf)
    if args.command == "check-summary":
        return run_check_summary(config, args.pdf)
    return run_once(config, args.pdf, force_regenerate=bool(args.force_regenerate))


if __name__ == "__main__":
    raise SystemExit(main())
